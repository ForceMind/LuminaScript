from fastapi import FastAPI, Depends, HTTPException, BackgroundTasks, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import delete, func, select, update
from sqlalchemy.orm import selectinload
from typing import List, Dict, Any, Literal, Optional
from pydantic import BaseModel, Field, field_validator
from contextlib import asynccontextmanager
import asyncio
import json
import re
from decimal import Decimal, ROUND_FLOOR

from database import get_db, SessionLocal
import models
import schemas
import auth
from services import llm  # Import LLM Service
import logging
import sys
from fastapi.responses import StreamingResponse
import io
from urllib.parse import quote

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

import database # needed for SessionLocal access in some scopes if not imported directly
from repositories.projects import (
    claim_generation,
    increment_tokens as increment_project_tokens,
    mark_claimed_failed as mark_claimed_project_failed,
    recover_interrupted,
)
from services.generation_state import (
    clear_generation_error,
    invalidate_scene_prompt_cache,
    record_generation_error,
)
from services.audit import log_ai_action, update_ai_action_status
from api.auth_routes import router as auth_router
from api.admin_routes import router as admin_router
from api.operations_routes import admin_router as admin_operations_router
from api.operations_routes import router as operations_router
from migrate import run_migrations
from services.admin_provisioning import ensure_admin_policy
from services.job_queue import CONTENT_JOB, OUTLINE_JOB, enqueue_job
from services.continuity import (
    build_content_continuity_context,
    build_outline_continuity_context,
    build_story_bible,
    looks_like_story_restart,
    summarize_scene_for_continuity,
)
from services.backups import backup_scheduler_loop
from services.project_access import (
    accessible_project_condition,
    project_role,
    require_project_access,
)
from services.prompt_templates import get_prompt_addendum
from services.usage import enforce_user_quota, invoke_with_quota
from services.versions import create_project_version
from services.setup_state import (
    assert_setup_writable,
    context_revision as setup_context_revision,
    detached_setup,
    revision_meta,
    valid_setup_cache,
    write_scene_prompt_cache,
    write_setup,
    write_setup_cache,
    write_working_draft,
)
from services import setup_fields
from services import setup_drafts

# Configure Logging
logging.basicConfig(
    level=logging.INFO, # Changed to INFO to avoid too much noise but capture essential flows
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("lumina_backend")

PROJECT_TYPE_LABELS = {
    "movie": "电影剧本",
    "tv": "剧集剧本",
    "short": "短剧剧本",
    "short_video": "短视频",
    "pending": "待确定"
}

SUMMARY_LABELS = {
    "title": "故事题目",
    "project_type": "剧本类型",
    "logline": "核心概念",
    "synopsis_brief": "故事梗概",
    "movie_duration": "电影时长",
    "scene_count_target": "目标场次",
    "episode_count": "集数",
    "episode_duration": "单集时长",
    "video_duration_seconds": "总时长",
    "tone": "基调",
    "time_period": "时代背景",
    "story_expansion": "剧情大纲",
    "character_details": "人物设定",
    "plot_details": "关键设定",
    "theme": "主题",
    "visual_style": "视觉风格",
    "user_notes": "补充说明"
}

SUMMARY_ORDER = [
    "title",
    "project_type",
    "logline",
    "synopsis_brief",
    "movie_duration",
    "scene_count_target",
    "episode_count",
    "episode_duration",
    "video_duration_seconds",
    "tone",
    "time_period",
    "story_expansion",
    "character_details",
    "plot_details",
    "theme",
    "visual_style",
    "user_notes"
]

SETUP_FLOW_STEPS = [
    {"key": "project_type", "question": "您想创作哪种类型的剧本？", "default_options": [
        {"label": "🎥 电影剧本", "value": "movie"},
        {"label": "📺 剧集剧本", "value": "tv"},
        {"label": "📱 短剧剧本", "value": "short"},
        {"label": "🎬 短视频", "value": "short_video"}
    ]},
    {"key": "movie_duration", "question": "电影预计时长是多少分钟？", "movie_only": True},
    {"key": "scene_count_target", "question": "您希望生成多少场戏？（电影通常 40-100 场，越多越细）", "movie_only": True},
    {"key": "episode_count", "question": "您计划创作多少集？", "tv_short_only": True},
    {"key": "episode_duration", "question": "每一集的大致时长是？", "tv_short_only": True},
    {"key": "video_duration_seconds", "question": "短视频总时长是多少秒？系统会自动按每 15 秒拆分。", "short_video_only": True},
    {"key": "tone", "question": "这部作品的基调是什么？"},
    {"key": "time_period", "question": "故事发生在什么时代背景？"},
    {"key": "story_expansion", "question": "我们需要基于目前构思扩展出完整的剧情大纲，您有什么特别想法吗？"},
    {"key": "character_details", "question": "主要角色的性格、外貌、关系或背景有什么特别设定？"},
    {"key": "plot_details", "question": "有哪些一定要发生的关键情节、转折或高潮？"},
    {"key": "title", "question": "现在请为这个故事确定一个题目，最好直接给出书名号里的名字。"},
    {"key": "theme", "question": "您想通过这个故事探讨什么主题？"},
    {"key": "visual_style", "question": "视觉风格偏向于什么？"},
    {"key": "user_notes", "question": "还有什么补充内容，或者特别要求吗？"},
    {"key": "final_confirm", "question": "以上是剧本的完整设定，请确认是否可以开始生成分场大纲？", "is_confirmation": True}
]

SETUP_MODE_KEY = "_setup_mode"
SETUP_MODE_FIELD = "setup_mode"
SETUP_MODE_AI_FAST = "ai_fast"
SETUP_MODE_GUIDED = "guided"
QUICK_REVIEW_FIELD = "quick_review"
QUICK_EDITED_FIELDS_KEY = "_quick_setup_user_edited_fields"
QUICK_CONTROL_FIELDS = {
    "project_type",
    "movie_duration",
    "scene_count_target",
    "episode_count",
    "episode_duration",
    "video_duration_seconds",
}

FINAL_CONFIRM_EDIT_TARGETS = [
    ("story_expansion", "返回修改剧情大纲"),
    ("character_details", "返回修改人物设定"),
    ("plot_details", "返回修改关键设定"),
    ("title", "返回修改故事题目"),
]

FINAL_CONFIRM_ALLOWED_VALUES = {"confirmed", "reset"} | {f"edit:{key}" for key, _ in FINAL_CONFIRM_EDIT_TARGETS}
AUTO_PREFILL_MIN_LENGTH = 120
AUTO_PREFILL_FLAG = "_auto_prefill_attempted"
AUTO_PREFILL_FIELDS = [
    # Keep basic setup questions mandatory.
    # We only auto-fill direction/content fields from long user input.
    "tone",
    "time_period",
    "title",
    "story_expansion",
    "character_details",
    "plot_details",
    "theme",
    "visual_style",
    "user_notes",
]

MAX_INTERACTION_ATTEMPTS = 2
MAX_INTERACTION_ANSWER_LENGTH = setup_fields.MAX_FIELD_LENGTH
MAX_QUICK_SETUP_TOTAL_LENGTH = setup_fields.MAX_TOTAL_LENGTH
ALLOWED_PROJECT_TYPES = setup_fields.PROJECT_TYPES
ALLOWED_INTERACTION_CONTEXT_KEYS = {
    step["key"] for step in SETUP_FLOW_STEPS
} | {SETUP_MODE_FIELD}
NUMERIC_INTERACTION_LIMITS = setup_fields.NUMERIC_LIMITS
GENERATION_TARGET_LIMITS = {
    "movie": 200,
    "tv": 100,
    "short": 100,
    "short_video": 40,
}
def get_relevant_setup_steps(project_type: str) -> List[Dict[str, Any]]:
    p_type = project_type or "movie"
    relevant_steps = []
    for step in SETUP_FLOW_STEPS:
        if step.get("movie_only") and p_type != "movie":
            continue
        if step.get("tv_short_only") and p_type not in {"tv", "short"}:
            continue
        if step.get("short_video_only") and p_type != "short_video":
            continue
        relevant_steps.append(step)
    return relevant_steps


def get_setup_value(
    project: models.Project,
    context: Dict[str, Any],
    key: str,
) -> Any:
    if key == "project_type":
        if project.project_type and project.project_type != "pending":
            return project.project_type
        return context.get(key)
    if key == "title":
        return context.get(key) or ""
    return context.get(key)


def has_existing_setup_progress(project: models.Project) -> bool:
    if project.project_type and project.project_type != "pending":
        return True
    context = build_normalized_context(project)
    return any(
        str(value or "").strip()
        for key, value in context.items()
        if key in {step["key"] for step in SETUP_FLOW_STEPS}
    )


def build_setup_context_revision(project: models.Project) -> str:
    return setup_context_revision(project)


def quick_setup_field_specs() -> List[Dict[str, Any]]:
    return [
        {"key": step["key"], "question": step["question"], "contract": setup_fields.field_contract(step["key"])}
        for step in SETUP_FLOW_STEPS
        if step["key"] != "final_confirm"
    ]


def normalize_quick_setup_values(
    project: models.Project,
    raw_values: Dict[str, Any],
    *,
    preserve_existing: bool = True,
) -> Dict[str, str]:
    current_context = build_normalized_context(project)
    setup_fields.validate_safety(raw_values)
    merged_values = dict(raw_values or {})
    if preserve_existing:
        for step in SETUP_FLOW_STEPS:
            key = step["key"]
            existing_value = get_setup_value(project, current_context, key)
            if str(existing_value or "").strip():
                merged_values[key] = existing_value

    merged_values.pop("final_confirm", None)
    return setup_fields.normalize_complete(merged_values)


def build_quick_review_sections(
    project: models.Project,
    values: Dict[str, str],
) -> List[Dict[str, Any]]:
    existing_context = build_normalized_context(project)
    sections: List[Dict[str, Any]] = []
    for step in get_relevant_setup_steps(values["project_type"]):
        key = step["key"]
        if key == "final_confirm":
            continue
        existing_value = get_setup_value(project, existing_context, key)
        sections.append(
            {
                "key": key,
                "label": SUMMARY_LABELS.get(key, key),
                "question": step["question"],
                "value": values[key],
                "editable": key not in QUICK_CONTROL_FIELDS,
                "source": "confirmed" if str(existing_value or "").strip() else "ai",
            }
        )
    return sections


def working_draft_payload(project: models.Project, draft: dict, *, stale: bool) -> Dict[str, Any]:
    values = draft.get("values") if isinstance(draft.get("values"), dict) else {}
    project_type = values.get("project_type")
    if project_type not in ALLOWED_PROJECT_TYPES:
        project_type = project.project_type if project.project_type in ALLOWED_PROJECT_TYPES else "movie"
    relevant = {step["key"] for step in get_relevant_setup_steps(project_type)} | set(values)
    sections = [{
        "key": step["key"], "label": SUMMARY_LABELS.get(step["key"], step["key"]),
        "question": step["question"], "value": str(values.get(step["key"], "")),
        "editable": not stale and step["key"] not in QUICK_CONTROL_FIELDS, "source": "ai",
    } for step in SETUP_FLOW_STEPS if step["key"] != "final_confirm" and step["key"] in relevant]
    return {
        "type": "interaction_required", "payload": {
            "field": QUICK_REVIEW_FIELD,
            "question": "保存的工作稿已过期，请只读查看、复制或明确丢弃。" if stale else "已恢复保存的快速工作稿，尚未写入正式设定。",
            "sections": sections, "values": values,
            "baseline_values": draft.get("baseline_values", {}),
            "edited_fields": draft.get("edited_fields", []),
            "ai_adjusted_fields": draft.get("ai_adjusted_fields", []),
            "draft_status": "stale" if stale else "saved", "draft_stale": stale,
            "read_only": stale, "base_setup_revision": draft.get("base_setup_revision"),
            "saved_at": draft.get("saved_at"),
        },
    }


def draft_response_state(project: models.Project) -> Dict[str, Any]:
    draft, stale = setup_drafts.inspect_draft(project)
    return {"quick_setup_draft": draft, "has_quick_setup_draft": draft is not None,
            "quick_setup_draft_stale": stale, "saved_draft_available": draft is not None,
            "draft_stale": stale}



def extract_story_title(raw_text: str) -> str:
    try:
        return setup_fields.normalize_field("title", raw_text)
    except ValueError:
        return ""


def sanitize_title_options(options: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    sanitized_options: List[Dict[str, Any]] = []
    seen_titles: set[str] = set()

    for option in options or []:
        if not isinstance(option, dict):
            continue

        raw_value = str(option.get("value", "") or "").strip()
        raw_label = str(option.get("label", "") or "").strip()
        clean_title = extract_story_title(raw_value)

        clean_title = clean_title.strip()
        if not clean_title or clean_title in seen_titles:
            continue

        seen_titles.add(clean_title)
        sanitized_options.append({
            "label": clean_title,
            "value": clean_title
        })

    return sanitized_options


def normalize_project_title(project: models.Project) -> bool:
    current_title = str(project.title or "").strip()
    context_title = ""
    if isinstance(project.global_context, dict):
        context_title = str(project.global_context.get("title", "") or "").strip()

    raw_candidate = context_title or current_title
    if not raw_candidate:
        return False

    clean_title = extract_story_title(raw_candidate)
    if not clean_title or clean_title == current_title:
        return False

    project.title = clean_title
    if isinstance(project.global_context, dict):
        updated_context = dict(project.global_context)
        updated_context["title"] = clean_title
        project.global_context = updated_context
    return True


def is_valid_character_details(value: Any) -> bool:
    try:
        setup_fields.normalize_field("character_details", value)
    except ValueError:
        return False
    return True


def normalize_project_context(project: models.Project) -> bool:
    if not isinstance(project.global_context, dict):
        return False

    updated_context = dict(project.global_context)
    changed = False

    if "character_details" in updated_context and not is_valid_character_details(updated_context.get("character_details")):
        updated_context.pop("character_details", None)
        changed = True

    if changed:
        project.global_context = updated_context

    return changed


def build_normalized_context(project: models.Project) -> Dict[str, Any]:
    raw_context = dict(project.global_context) if isinstance(project.global_context, dict) else {}
    context = {
        key: value
        for key, value in raw_context.items()
        if not str(key).startswith("_")
    }
    if project.project_type and project.project_type != "pending":
        context["project_type"] = project.project_type
    return context


def get_internal_project_context(project: models.Project) -> Dict[str, Any]:
    return dict(project.global_context) if isinstance(project.global_context, dict) else {}


def has_setup_value(project: models.Project, context: Dict[str, Any], key: str) -> bool:
    if key == "project_type":
        value = project.project_type if project.project_type and project.project_type != "pending" else context.get("project_type")
    elif key == "title":
        value = context.get("title") or project.title
    else:
        value = context.get(key)

    if isinstance(value, str):
        value = value.strip()

    return value not in (None, "", "pending")


def normalize_extracted_setup_value(key: str, value: Any) -> str:
    try:
        return setup_fields.normalize_field(key, value)
    except ValueError:
        return ""


def validate_interaction_answer(context_key: str, raw_answer: Any) -> str:
    if context_key in setup_fields.SETUP_FIELDS:
        try:
            return setup_fields.normalize_field(context_key, raw_answer)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
    answer = str(raw_answer or "").strip()
    if not answer:
        raise HTTPException(status_code=422, detail="回答不能为空")
    if len(answer) > MAX_INTERACTION_ANSWER_LENGTH:
        raise HTTPException(
            status_code=422,
            detail=f"回答不能超过 {MAX_INTERACTION_ANSWER_LENGTH} 个字符",
        )


    return answer


def should_auto_prefill_from_logline(project: models.Project, context: Dict[str, Any]) -> bool:
    if not isinstance(context, dict):
        return False
    raw_context = get_internal_project_context(project)
    if raw_context.get(AUTO_PREFILL_FLAG):
        return False

    clean_logline = re.sub(r"\s+", "", str(project.logline or ""))
    return len(clean_logline) >= AUTO_PREFILL_MIN_LENGTH


def apply_auto_prefill(project: models.Project, extracted_payload: Dict[str, Any] | None) -> tuple[List[str], bool]:
    current_context = dict(project.global_context) if isinstance(project.global_context, dict) else {}
    extracted_payload = extracted_payload if isinstance(extracted_payload, dict) else {}
    setup_fields.validate_safety(extracted_payload)
    changed = False
    filled_fields: List[str] = []

    if not current_context.get(AUTO_PREFILL_FLAG):
        current_context[AUTO_PREFILL_FLAG] = True
        changed = True

    for key in AUTO_PREFILL_FIELDS:
        if has_setup_value(project, current_context, key):
            continue

        if key == "title":
            raw_value = extracted_payload.get(key) or project.logline or ""
        else:
            raw_value = extracted_payload.get(key)

        normalized_value = normalize_extracted_setup_value(key, raw_value)
        if not normalized_value:
            continue
        try:
            setup_fields.validate_safety({
                **{field: value for field, value in current_context.items() if field in setup_fields.SETUP_FIELDS},
                key: normalized_value,
            })
        except ValueError:
            continue

        if key == "project_type":
            project.project_type = normalized_value

        if key == "title":
            project.title = normalized_value

        current_context[key] = normalized_value
        filled_fields.append(key)
        changed = True

    if changed:
        project.global_context = current_context

    return filled_fields, changed


def should_invalidate_cached_question(cache_payload: Any, current_context: Dict[str, Any] | None = None) -> bool:
    if not isinstance(cache_payload, dict):
        return False

    payload = cache_payload.get("payload")
    if not isinstance(payload, dict):
        return False

    field = payload.get("field")
    options = payload.get("options")
    current_context = current_context or {}

    if field == "project_type":
        current_value = current_context.get("project_type")
        if isinstance(current_value, str):
            current_value = current_value.strip()
        if current_value not in (None, "", "pending"):
            return True

    if field == "retry_current_step":
        return True

    if field and field not in {"final_confirm", "project_type"}:
        current_value = current_context.get(field)
        if isinstance(current_value, str):
            current_value = current_value.strip()
        if current_value not in (None, ""):
            return True

    if field == "final_confirm":
        if not isinstance(options, list):
            return True
        if not any(isinstance(option, dict) and str(option.get("value", "")).startswith("edit:") for option in options):
            return True

    if field == "title":
        if any(key not in current_context for key in ("story_expansion", "character_details", "plot_details")):
            return True

    if field not in {"character_details", "story_expansion", "plot_details"}:
        return False
    if not isinstance(options, list) or len(options) < 3:
        return True

    for option in options:
        if not isinstance(option, dict):
            return True
        try:
            setup_fields.normalize_field(field, option.get("value"))
        except ValueError:
            return True
    return False


def has_valid_setup_cache(project: models.Project, *, mode: str, stage: str) -> bool:
    if not valid_setup_cache(project, mode=mode, stage=stage):
        return False
    if stage == QUICK_REVIEW_FIELD:
        sections = project.next_step_cache["payload"]["sections"]
        values = {item["key"]: item["value"] for item in sections}
        if len(values) != len(sections):
            return False
        try:
            normalized = normalize_quick_setup_values(project, values, preserve_existing=False)
        except ValueError:
            return False
        if normalized != values:
            return False
    return not should_invalidate_cached_question(project.next_step_cache, build_normalized_context(project))


def rewind_project_setup(project: models.Project, target_key: str) -> Dict[str, Any]:
    current_context = dict(project.global_context) if isinstance(project.global_context, dict) else {}
    project_type = project.project_type if project.project_type and project.project_type != "pending" else current_context.get("project_type", "movie")
    relevant_steps = get_relevant_setup_steps(project_type)

    clear_from_here = False
    for step in relevant_steps:
        key = step["key"]
        if key == target_key:
            clear_from_here = True
        if not clear_from_here or key == "project_type":
            continue

        current_context.pop(key, None)
        if key == "title":
            project.title = ""

    for derived_key in (
        "synopsis_brief",
        "synopsis_detailed",
        "brief_synopsis",
        "detailed_synopsis",
        "story_brief",
        "story_detailed",
        "final_confirm",
    ):
        current_context.pop(derived_key, None)

    project.global_context = current_context
    project.next_step_cache = None
    return current_context


def format_summary_value(key: str, value: Any) -> str:
    if value is None:
        return ""

    if isinstance(value, (list, dict)):
        text = json.dumps(value, ensure_ascii=False, indent=2)
    else:
        text = str(value).strip()

    if not text:
        return ""

    if key == "project_type":
        return PROJECT_TYPE_LABELS.get(text, text)

    if key in {"movie_duration", "episode_duration"}:
        try:
            return f"{setup_fields.normalize_number(key, text).removesuffix('mins')} 分钟"
        except ValueError:
            return text

    if key == "video_duration_seconds":
        try:
            return f"{setup_fields.normalize_number(key, text)} 秒"
        except ValueError:
            return text

    if key == "scene_count_target" and re.fullmatch(r"\d+", text):
        return f"{text} 场"

    if key == "episode_count" and re.fullmatch(r"\d+", text):
        return f"{text} 集"

    return text


def build_context_summary(project: models.Project, context: Dict[str, Any]) -> str:
    summary_context = dict(context or {})
    if project.logline:
        summary_context.setdefault("logline", project.logline)

    lines: List[str] = []
    for key in SUMMARY_ORDER:
        if key not in summary_context:
            continue

        label = SUMMARY_LABELS.get(key)
        if not label:
            continue

        display_value = format_summary_value(key, summary_context.get(key))
        if not display_value:
            continue

        if "\n" in display_value:
            lines.append(f"- {label}：")
            lines.append(display_value)
        else:
            lines.append(f"- {label}：{display_value}")

    return "\n".join(lines)


async def ensure_story_synopsis(
    project: models.Project,
    context: Dict[str, Any],
    db: Optional[AsyncSession] = None,
    *,
    persist: bool = True,
    actor_id: Optional[int] = None,
    optional: bool = False,
) -> Dict[str, Any]:
    enriched_context = dict(context or {})
    has_brief = bool(str(enriched_context.get("synopsis_brief", "") or "").strip())
    has_detailed = bool(str(enriched_context.get("synopsis_detailed", "") or "").strip())
    if has_brief and has_detailed:
        return enriched_context

    if db is None:
        raise RuntimeError("Synopsis generation requires a database session for quota and audit")
    synopsis_revision = build_setup_context_revision(project)

    async def synopsis_is_stale():
        return build_setup_context_revision(project) != synopsis_revision

    if optional:
        try:
            await enforce_user_quota(db, project.owner_id)
        except HTTPException as exc:
            if exc.status_code == 429:
                return enriched_context
            raise
    try:
        synopsis, _ = await run_project_ai_call(
            db=db, project=project, actor_id=actor_id or project.owner_id,
            action="generate_story_synopsis", prompt=json.dumps(enriched_context, ensure_ascii=False),
            stale_check=synopsis_is_stale,
            invoke=lambda: llm.generate_story_synopsis(
                logline=project.logline or "", context=enriched_context,
                project_type=project.project_type or "movie"
            ),
        )
    except HTTPException as exc:
        if optional and exc.status_code == 429:
            return enriched_context
        raise
    except Exception as exc:
        raise HTTPException(status_code=503, detail="AI 梗概生成失败，原设定未改变，请重试。") from exc

    brief = str(synopsis.get("brief", "") or "").strip()
    detailed = str(synopsis.get("detailed", "") or "").strip()

    if brief:
        enriched_context["synopsis_brief"] = brief
    if detailed:
        enriched_context["synopsis_detailed"] = detailed

    if persist and enriched_context != (project.global_context or {}):
        project.global_context = enriched_context

    return enriched_context

async def recover_interrupted_generations() -> None:
    """Mark in-process jobs interrupted by a previous process exit as failed."""
    async with SessionLocal() as db:
        recovered_scenes, recovered_projects = await recover_interrupted(db)
        if recovered_scenes or recovered_projects:
            logger.warning(
                "Recovered interrupted generation state: projects=%s scenes=%s",
                recovered_projects,
                recovered_scenes,
            )


@asynccontextmanager
async def lifespan(_: FastAPI):
    logger.info("服务器正在启动...")

    try:
        logger.info("Running database migrations...")
        await asyncio.to_thread(run_migrations)
        logger.info("Database migrations complete.")

        logger.info("Running administrator policy check...")
        administrators = await ensure_admin_policy()
        logger.info(
            "Administrator policy check complete. Administrators=%s",
            administrators,
        )
    except Exception as e:
        logger.exception(f"Failed to run schema upgrade: {e}")
        raise

    await recover_interrupted_generations()
    logger.info("数据库初始化完成，服务准备就绪。")
    backup_task = asyncio.create_task(backup_scheduler_loop())
    try:
        yield
    finally:
        backup_task.cancel()
        try:
            await backup_task
        except asyncio.CancelledError:
            pass


# Initialize App after the startup helpers are defined.
app = FastAPI(title="LuminaScript API", version="0.1.0", lifespan=lifespan)
app.include_router(auth_router)
app.include_router(admin_router)
app.include_router(operations_router)
app.include_router(admin_operations_router)

@app.get("/")
async def root():
    logger.info("收到根路径请求")
    return {"message": "欢迎使用妙笔流光 (LuminaScript) API"}

# --- Project Management ---

@app.post("/projects/", response_model=schemas.ProjectResponse)
async def create_project(
    project: schemas.ProjectCreate, 
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    logger.info(f"用户 {current_user.username} 正在创建新项目，Logline: {project.logline[:50]}...")
    # 1. First step: Create the project record based on logline
    # Real implementation would call LLM here to analyze logline first, 
    # but for now we just save it.
    new_project = models.Project(
        title=project.title,
        logline=project.logline,
        project_type=project.project_type,
        owner_id=current_user.id
    )
    db.add(new_project)
    await db.commit()
    await db.refresh(new_project)
    
    logger.info(f"项目创建成功 ID: {new_project.id}")

    # Reload to ensure relationships (scenes) are loaded for Pydantic
    result = await db.execute(
        select(models.Project)
        .where(models.Project.id == new_project.id)
        .options(selectinload(models.Project.scenes))
    )
    return result.scalars().first()


@app.get("/projects/", response_model=List[schemas.ProjectListResponse])
async def list_projects(
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    result = await db.execute(
        select(models.Project)
        .where(accessible_project_condition(current_user.id))
        .order_by(models.Project.id.desc())
    )
    projects = result.scalars().all()
    for project in projects:
        project.access_role = await project_role(db, project, current_user.id) or "viewer"

    return projects


@app.get("/projects/{project_id}", response_model=schemas.ProjectResponse)
async def get_project(
    project_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    project, role = await require_project_access(
        db,
        project_id,
        current_user.id,
        load_scenes=True,
    )
    project.access_role = role

    return project


@app.delete("/projects/{project_id}")
async def delete_project(
    project_id: int, 
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    project, role = await require_project_access(db, project_id, current_user.id)
    if role != "owner":
        raise HTTPException(status_code=403, detail="只有项目所有者可以删除项目")

    # Mark as failed/deleted to stop background tasks
    project.status = models.ProcessingStatus.FAILED 
    await db.delete(project)
    await db.commit()
    return {"status": "success"}

@app.patch("/projects/{project_id}", response_model=schemas.ProjectResponse)
async def update_project(
    project_id: int,
    project_update: schemas.ProjectUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    # Use select with options to eager load scenes to avoid MissingGreenlet error in response validation
    result = await db.execute(
        select(models.Project)
        .where(models.Project.id == project_id)
        .options(selectinload(models.Project.scenes))
    )
    project = result.scalars().first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    role = await project_role(db, project, current_user.id)
    if role not in {"owner", "editor"}:
        raise HTTPException(status_code=404, detail="Project not found")
    project.access_role = role

    await assert_setup_writable(db, project, project_update.context_revision)
    if project_update.project_type:
        current_context = get_internal_project_context(project)
        current_context.pop("final_confirm", None)
        current_context["project_type"] = project_update.project_type
        await write_setup(db, project, project_update.context_revision, {
            "project_type": project_update.project_type,
            "global_context": current_context,
        })
    await db.commit()
    return project

class InteractionRequest(BaseModel):
    answer: str = Field(max_length=MAX_INTERACTION_ANSWER_LENGTH)
    context_key: str = Field(min_length=1, max_length=100)
    context_revision: Optional[str] = Field(default=None, max_length=128)

    @field_validator("context_key")
    @classmethod
    def validate_context_key(cls, value: str) -> str:
        normalized = value.strip()
        if normalized not in ALLOWED_INTERACTION_CONTEXT_KEYS:
            raise ValueError("不支持的交互字段")
        return normalized


class QuickSetupReviewRequest(BaseModel):
    action: Literal["confirm", "guided", "save", "save_guided", "discard", "regenerate"] = "confirm"
    values: Dict[str, str] = Field(default_factory=dict)
    baseline_values: Dict[str, str] = Field(default_factory=dict)
    edited_fields: List[str] = Field(default_factory=list, max_length=20)
    ai_adjusted_fields: List[str] = Field(default_factory=list, max_length=20)
    context_revision: Optional[str] = Field(default=None, max_length=128)

    @field_validator("values", "baseline_values")
    @classmethod
    def validate_quick_setup_values(cls, value: Dict[str, str]) -> Dict[str, str]:
        return setup_fields.validate_safety(value)

    @field_validator("edited_fields", "ai_adjusted_fields")
    @classmethod
    def validate_edited_fields(cls, value: List[str]) -> List[str]:
        allowed = {step["key"] for step in SETUP_FLOW_STEPS if step["key"] != "final_confirm"}
        return list(dict.fromkeys(str(item).strip() for item in value if str(item).strip() in allowed))


class QuickSetupAIReviseRequest(BaseModel):
    operation: Literal["regenerate_field", "review_edits"]
    scope: Literal["edited_only", "related"] = "edited_only"
    values: Dict[str, str] = Field(default_factory=dict)
    baseline_values: Dict[str, str] = Field(default_factory=dict)
    target_field: Optional[str] = None
    edited_fields: List[str] = Field(default_factory=list, max_length=20)
    ai_adjusted_fields: List[str] = Field(default_factory=list, max_length=20)
    context_revision: Optional[str] = Field(default=None, max_length=128)
    instruction: Optional[str] = Field(default=None, max_length=MAX_INTERACTION_ANSWER_LENGTH)

    @field_validator("values", "baseline_values")
    @classmethod
    def validate_values(cls, value: Dict[str, str]) -> Dict[str, str]:
        return setup_fields.validate_safety(value)

    @field_validator("target_field", "instruction")
    @classmethod
    def strip_optional(cls, value: Optional[str]) -> Optional[str]:
        return value.strip() if isinstance(value, str) else value

    @field_validator("edited_fields", "ai_adjusted_fields")
    @classmethod
    def validate_revision_fields(cls, value: List[str]) -> List[str]:
        allowed = {
            step["key"]
            for step in SETUP_FLOW_STEPS
            if step["key"] != "final_confirm"
        }
        result = list(
            dict.fromkeys(str(key).strip() for key in value if str(key).strip())
        )
        invalid = [k for k in result if k not in allowed]
        if invalid:
            raise ValueError(f"不可编辑字段: {', '.join(invalid)}")
        return result


async def write_setup_ai_audit(**entry: Any) -> Optional[int]:
    """Fail closed before exposing/applying setup AI results if audit is down."""
    try:
        return await log_ai_action(**entry)
    except Exception as exc:
        logger.error("Setup AI audit write failed: %s", type(exc).__name__)
        raise HTTPException(
            status_code=503,
            detail="AI 审计写入失败，原始返回记录未能确认保存，本次 AI 结果未应用；项目存在时已保留已知项目 Token，账户用量记录可能缺失，请管理员检查后重试。",
        ) from exc


async def run_project_ai_call(
    *, db: AsyncSession, project: models.Project, actor_id: int,
    action: str, prompt: str, invoke: Any, attempt: int = 1,
    expected_status: Optional[models.ProcessingStatus] = None,
    validate: Any = None,
    stale_check: Any = None,
    result_status: Any = None,
    audit_ids: Optional[List[int]] = None,
) -> tuple[Any, int]:
    """Account every real call before applying its output; capture billing before await.

    Quotas are soft checks, not reservations: concurrent in-flight requests may exceed
    a limit. Known usage survives invalid responses, cancellation and audit failure.
    """
    project_id, billed_user_id = project.id, project.owner_id
    await enforce_user_quota(db, billed_user_id)
    result, usage, error, raw = None, 0, None, ""
    try:
        result, usage = await invoke_with_quota(billed_user_id, invoke)
        raw = getattr(result, "raw_content", None)
        if raw is None:
            raw = result if isinstance(result, str) else json.dumps(result, ensure_ascii=False)
        if validate is not None:
            validate(result)
    except Exception as exc:
        if isinstance(exc, HTTPException) and exc.status_code == 429:
            raise
        error = exc
        usage = max(int(usage or 0), int(getattr(exc, "usage", 0) or 0))
        raw = str(getattr(exc, "raw_content", "") or raw)
    usage = max(0, int(usage or 0))
    if usage:
        await increment_project_tokens(db, project, usage)
        await db.commit()
    live_project = await db.scalar(select(models.Project).where(models.Project.id == project_id).execution_options(populate_existing=True))
    stale = live_project is None or (expected_status is not None and live_project.status != expected_status)
    if not stale and stale_check is not None:
        stale = await stale_check()
    audit_id = await write_setup_ai_audit(
        user_id=actor_id, billed_user_id=billed_user_id,
        project_id=project_id if live_project is not None else None,
        action=action, prompt=prompt, response=raw, tokens=usage,
        status="stale" if stale else "failed" if error else result_status(result) if result_status else "success", attempt=attempt,
        error_type="stale_context" if stale else getattr(error, "error_type", type(error).__name__) if error else None,
        error_message="AI 返回前项目已删除或生成已停止" if stale else str(error) if error else None,
    )
    if audit_ids is not None and audit_id is not None:
        audit_ids.append(audit_id)
    if stale:
        raise HTTPException(status_code=409, detail="AI 返回前项目已变化，本次结果未应用。")
    if error:
        raise error
    return result, usage


def validate_ai_text(value: Any) -> None:
    if not str(value or "").strip():
        raise ValueError("AI 返回了空正文")


def validate_outline_batch(value: Any) -> None:
    if not isinstance(value, list) or len(value) != 1:
        raise ValueError("AI 分场大纲必须返回一个场次")
    if not isinstance(value[0], dict) or not str(value[0].get("outline") or "").strip():
        raise ValueError("AI 返回了空分场大纲")


def validate_continuation(value: Any, scene_index: int, *, outline: bool = False) -> None:
    if outline:
        validate_outline_batch(value)
        text = value[0]["outline"]
    else:
        validate_ai_text(value)
        text = str(value)
    if looks_like_story_restart(text, scene_index):
        raise ValueError("连续性守卫拒绝了疑似重新开篇的内容")


async def record_quick_setup_ai_revision(
    *,
    db: AsyncSession,
    project: models.Project,
    current_user: models.User,
    operation: str,
    prompt: str,
    response: str,
    tokens: int,
    status: str,
    billed_user_id: Optional[int] = None,
    error_type: Optional[str] = None,
    error_message: Optional[str] = None,
) -> None:
    """Charge known usage and write an audit entry without exposing provider errors."""
    token_count = max(0, int(tokens or 0))
    if token_count:
        await increment_project_tokens(db, project, token_count)
        await db.commit()

    action = (
        "regenerate_quick_setup_field"
        if operation == "regenerate_field"
        else "review_quick_setup_edits"
    )
    await write_setup_ai_audit(
        user_id=current_user.id, project_id=project.id, action=action,
        billed_user_id=billed_user_id if billed_user_id is not None else project.owner_id,
        prompt=prompt, response=response, tokens=token_count, status=status,
        step_key=QUICK_REVIEW_FIELD, error_type=error_type, error_message=error_message,
    )


class ContentReviewRequest(BaseModel):
    text: str = Field(min_length=1, max_length=MAX_INTERACTION_ANSWER_LENGTH)


@app.post("/content/review")
async def review_content(
    payload: ContentReviewRequest,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    """
    Review free user text and return whether it should be rewritten,
    plus an AI-generated safe rewrite suggestion.
    """
    if llm.review_requires_ai(payload.text):
        await enforce_user_quota(db, current_user.id)
    try:
        template_instructions = await get_prompt_addendum(
            db,
            stage="review",
            project_type="all",
        )
        result = await invoke_with_quota(current_user.id, lambda: llm.review_user_input(
            payload.text,
            template_instructions=template_instructions,
        ))
    except Exception as e:
        if isinstance(e, HTTPException) and e.status_code == 429:
            raise
        await write_setup_ai_audit(
            user_id=current_user.id, billed_user_id=current_user.id, project_id=None,
            action="review_content", prompt=payload.text,
            response=str(getattr(e, "raw_content", "") or ""), tokens=max(0, int(getattr(e, "usage", 0) or 0)),
            status="failed", error_type=getattr(e, "error_type", type(e).__name__), error_message=str(e),
        )
        raise HTTPException(status_code=503, detail="Content review service unavailable") from e
    usage = max(0, int(getattr(result, "usage", 0) or 0))
    if getattr(result, "ai_called", llm.review_requires_ai(payload.text)):
        await write_setup_ai_audit(
            user_id=current_user.id, billed_user_id=current_user.id, project_id=None,
            action="review_content", prompt=payload.text,
            response=getattr(result, "raw_content", json.dumps(result, ensure_ascii=False)), tokens=usage,
        )
    return {**result, "tokens_used": usage}

@app.post("/projects/{project_id}/interact")
async def submit_interaction(
    project_id: int,
    interaction: InteractionRequest,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    logger.info(f"收到项目 {project_id} 的交互回答: Key={interaction.context_key}, Answer={interaction.answer}")
    
    project, _ = await require_project_access(
        db,
        project_id,
        current_user.id,
        minimum_role="editor",
    )

    await assert_setup_writable(db, project, interaction.context_revision)
    draft = detached_setup(project)
    current_context = get_internal_project_context(draft)
    previous_title = project.title
    
    # Special Handling: Reset
    answer_text = validate_interaction_answer(
        interaction.context_key,
        interaction.answer,
    )
    if interaction.context_key == SETUP_MODE_FIELD:
        if answer_text not in {SETUP_MODE_AI_FAST, SETUP_MODE_GUIDED}:
            raise HTTPException(status_code=422, detail="不支持的设定方式")
        current_context[SETUP_MODE_KEY] = answer_text
        await write_setup(db, project, interaction.context_revision, {
            "global_context": current_context,
            "quick_setup_draft": setup_drafts.draft_after_mode_change(project),
        })
        await db.commit()
        return {
            "status": "setup_mode_updated",
            "setup_mode": answer_text,
            "context": project.global_context,
            "title": project.title or previous_title or "",
            "total_tokens": int(project.total_tokens or 0),
            **revision_meta(project),
            **draft_response_state(project),
        }

    if interaction.context_key == 'final_confirm' and answer_text == 'reset':
        logger.info(f"项目 {project_id} 收到重置请求，清空上下文重新开始设定流程")
        await write_setup(db, project, interaction.context_revision, {
            "global_context": {}, "project_type": "pending",
        })
        await db.commit()
        return {"status": "reset", "context": {}, **revision_meta(project)}

    if interaction.context_key == 'final_confirm' and answer_text.startswith('edit:'):
        target_key = answer_text.split(':', 1)[1].strip()
        rewind_project_setup(draft, target_key)
        await write_setup(db, project, interaction.context_revision, {
            "global_context": draft.global_context, "title": draft.title,
        })
        await db.commit()
        return {
            "status": "rewind",
            "context": project.global_context,
            "title": project.title or previous_title or "",
            **revision_meta(project),
        }

    if interaction.context_key == 'final_confirm' and answer_text not in FINAL_CONFIRM_ALLOWED_VALUES:
        raise HTTPException(status_code=400, detail="请直接点击下方按钮选择确认操作，再重新发起。")

    # Ensure project_type is synced if that was the key (legacy support)
    if interaction.context_key == 'project_type':
        draft.project_type = answer_text
    
    # Handle Title Update specifically
    if interaction.context_key == 'title':
        logger.info(f"Checking title update. Proposed Title: '{answer_text}'")
        clean_title = extract_story_title(answer_text)
        if clean_title:
            current_context[interaction.context_key] = clean_title
            draft.title = clean_title
            logger.info(f"Project Title Updated to: {draft.title}")
        else:
             current_context.pop(interaction.context_key, None)
             logger.warning(f"Ignored suspicious title update: {answer_text}")
    else:
        current_context[interaction.context_key] = answer_text

    if interaction.context_key != "final_confirm":
        current_context.pop("final_confirm", None)
    if interaction.context_key == "final_confirm" and answer_text == "confirmed":
        try:
            normalized = setup_fields.normalize_complete({
                **{key: value for key, value in current_context.items() if key in setup_fields.SETUP_FIELDS},
                "project_type": draft.project_type,
            })
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        current_context.update(normalized)
        draft.title = normalized["title"]
    try:
        setup_fields.validate_safety({key: value for key, value in current_context.items() if key in setup_fields.SETUP_FIELDS})
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    await write_setup(db, project, interaction.context_revision, {
        "global_context": current_context,
        "project_type": draft.project_type,
        "title": draft.title,
    })

    await db.commit()
    logger.info(f"项目 {project_id} 上下文已更新，缓存已清除")
    return {
        "status": "updated",
        "context": project.global_context,
        "title": project.title or previous_title or "",
        "total_tokens": int(project.total_tokens or 0),
        **revision_meta(project),
    }


async def generate_validated_setup_options(
    *, db: AsyncSession, project: models.Project, current_user: models.User,
    step_key: str, question: str, values: Dict[str, str], context: str,
    revision: str, template_instructions: str, action: str,
) -> tuple[Dict[str, Any], int]:
    """Keep good options and fill only the deficit, with at most two AI calls."""
    setup_fields.validate_safety(values)
    accepted: List[Dict[str, str]] = []
    previous_value = ""
    if step_key in values and values[step_key].strip():
        try:
            previous_value = setup_fields.normalize_field(step_key, values[step_key])
        except ValueError:
            previous_value = values[step_key].strip()
    excluded = {previous_value} if previous_value else set()
    total_usage = 0
    billed_user_id = project.owner_id
    final_question = question
    rejection_summary: List[Dict[str, Any]] = []
    for attempt in range(1, MAX_INTERACTION_ATTEMPTS + 1):
        await enforce_user_quota(db, billed_user_id)
        rejected: Dict[str, int] = {}

        def reject(reason: str, count: int = 1) -> None:
            bounded = str(reason)[:160]
            if bounded not in rejected and len(rejected) >= 8:
                bounded = "其他格式或内容问题"
                if bounded not in rejected:
                    return
            rejected[bounded] = rejected.get(bounded, 0) + max(1, int(count))

        call_context = json.dumps({
            "context": context, "current_draft": values,
            "target_field": step_key, "required_count": 3 - len(accepted),
            "excluded_values": sorted(excluded),
            "accepted_options": accepted,
            "field_contract": setup_fields.field_contract(step_key),
            "available_value_characters": min(MAX_INTERACTION_ANSWER_LENGTH, MAX_QUICK_SETUP_TOTAL_LENGTH - sum(len(value) for key, value in values.items() if key != step_key)),
            "rejection_summary": rejection_summary,
            "rule": "只补足缺额；不得重复排除值。其他字段原文锁定，value必须是完整有效内容。",
        }, ensure_ascii=False)
        usage = 0
        raw_content = ""
        error_type = None
        error_message = None
        try:
            data, usage = await invoke_with_quota(billed_user_id, lambda: llm.generate_interaction_options(
                step_key, question, call_context,
                template_instructions=template_instructions,
            ))
            raw_content = getattr(data, "raw_content", json.dumps(data, ensure_ascii=False))
            for item in getattr(data, "rejection_summary", [])[:8]:
                reject(item.get("reason", "字段格式不符合约束"), item.get("count", 1))
            if not isinstance(data, dict) or not isinstance(data.get("options"), list):
                raise ValueError("AI 选项不是有效列表")
            proposed_question = data.get("question")
            if isinstance(proposed_question, str) and 0 < len(proposed_question) <= MAX_INTERACTION_ANSWER_LENGTH:
                final_question = proposed_question
            for option in data["options"]:
                if not isinstance(option, dict):
                    reject("选项不是对象")
                    continue
                label = option.get("label")
                if not isinstance(label, str) or not label.strip() or len(label) > MAX_INTERACTION_ANSWER_LENGTH:
                    reject("label必须为非空且未超长的文本")
                    continue
                if not isinstance(option.get("value"), str):
                    reject("value必须为文本，不能是对象、列表或缺失值")
                    continue
                try:
                    normalized = setup_fields.normalize_field(step_key, option.get("value"))
                    setup_fields.validate_safety({**values, step_key: normalized})
                except ValueError as exc:
                    reject(str(exc))
                    continue
                if normalized in excluded:
                    reject("规范值与当前旧值相同" if normalized == previous_value else "规范值与已保留选项重复")
                    continue
                accepted.append({"label": label.strip(), "value": normalized})
                excluded.add(normalized)
                if len(accepted) == 3:
                    break
        except Exception as exc:
            if isinstance(exc, HTTPException) and exc.status_code == 429:
                raise
            usage = max(int(usage or 0), int(getattr(exc, "usage", 0) or 0))
            raw_content = str(getattr(exc, "raw_content", "") or raw_content)
            error_type = str(getattr(exc, "error_type", type(exc).__name__))
            error_message = str(exc)
            reject(f"{error_type}: {error_message}")
        rejection_summary = [{"reason": reason, "count": count} for reason, count in rejected.items()]
        usage = max(0, int(usage or 0))
        total_usage += usage
        if usage:
            await increment_project_tokens(db, project, usage)
            await db.commit()
        stale = False
        try:
            await assert_setup_writable(db, project, revision)
            if revision != build_setup_context_revision(project):
                stale = True
        except HTTPException as exc:
            if exc.status_code != 409:
                raise
            stale = True
        status = "stale" if stale else "success" if len(accepted) == 3 else "partial" if accepted and attempt < MAX_INTERACTION_ATTEMPTS else "failed"
        await write_setup_ai_audit(
            user_id=current_user.id, project_id=project.id, action=action,
            billed_user_id=billed_user_id,
            prompt=call_context, response=raw_content, tokens=usage, status=status,
            step_key=QUICK_REVIEW_FIELD if action == "regenerate_quick_setup_field" else step_key,
            error_type="stale_context" if stale else error_type or (None if len(accepted) == 3 else "insufficient_valid_options"),
            error_message="AI 返回前设定已更新或开始生成" if stale else error_message,
            attempt=attempt,
        )
        if stale:
            raise HTTPException(status_code=409, detail="AI 分析期间项目设定已更新，请刷新后重试。")
        if len(accepted) == 3:
            return {"question": final_question, "options": accepted}, total_usage
    raise HTTPException(status_code=503, detail=f"AI 补齐后仍仅有 {len(accepted)} 个有效新选项，需要 3 个；原草案未改变，请重试。")


@app.post("/projects/{project_id}/setup/quick-review/ai-revise")
async def revise_quick_setup_with_ai(
    project_id: int, payload: QuickSetupAIReviseRequest,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user),
):
    project, _ = await require_project_access(db, project_id, current_user.id, minimum_role="editor")
    await assert_setup_writable(db, project, payload.context_revision)
    setup_drafts.require_current_draft(project)
    if payload.context_revision != build_setup_context_revision(project):
        raise HTTPException(status_code=409, detail="项目设定已更新，请刷新后重试。")
    # Only safety/shape is global here. Invalid repair targets must reach AI;
    # unrelated values (including whitespace and non-canonical units) stay exact.
    try:
        values = setup_fields.validate_safety(payload.values)
        all_fields = setup_fields.relevant_fields(values.get("project_type", project.project_type or ""))
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    baseline, baseline_source = setup_drafts.baseline_for(project, payload.baseline_values)
    actual_changes = setup_drafts.value_changes(values, baseline)
    saved_metadata, _ = setup_drafts.inspect_draft(project)
    ai_sources = payload.ai_adjusted_fields if "ai_adjusted_fields" in payload.model_fields_set else (saved_metadata or {}).get("ai_adjusted_fields", [])
    manual_sources = payload.edited_fields if "edited_fields" in payload.model_fields_set else (saved_metadata or {}).get("edited_fields", [])
    changed_fields = {key: {**change, "source": "ai" if key in ai_sources and key not in manual_sources else "manual"}
                      for key, change in actual_changes.items()}
    locked_fields: List[str] = []
    invalid_changed_fields: Dict[str, str] = {}
    if payload.operation == "regenerate_field":
        if not payload.target_field or payload.target_field not in all_fields or payload.target_field == "project_type":
            raise HTTPException(status_code=422, detail="单项重生必须指定有效 target_field")
        allowed = [payload.target_field]
    else:
        allowed = [field for field in all_fields if field in actual_changes and field not in QUICK_CONTROL_FIELDS]
        if payload.scope == "related":
            locked_fields = sorted(set(actual_changes) | QUICK_CONTROL_FIELDS)
            allowed = [field for field in all_fields if field not in locked_fields]
            for field, change in actual_changes.items():
                try:
                    setup_fields.normalize_field(field, change["after"])
                except ValueError as exc:
                    invalid_changed_fields[field] = str(exc)
        else:
            locked_fields = sorted(set(all_fields) - set(allowed))
        if not actual_changes or not allowed:
            note = "当前值与基线一致，无需额外 AI 分析。" if not actual_changes else "当前改后值已锁定，没有可调整的关联内容项。"
            if invalid_changed_fields:
                note += " 锁定改后字段仍需人工修正：" + "、".join(invalid_changed_fields)
            return {"status": "candidate", "operation": payload.operation, "scope": payload.scope,
                    "changes": [], "changed_fields": [], "summary": note, "tokens_used": 0,
                    "total_tokens": int(project.total_tokens or 0), "context_revision": payload.context_revision}
    await enforce_user_quota(db, project.owner_id)
    billed_user_id = project.owner_id
    template_instructions = await get_prompt_addendum(db, stage="interaction", project_type=values.get("project_type", "all"))
    audit_prompt = json.dumps({
        "operation": payload.operation, "scope": payload.scope,
        "target_field": payload.target_field, "edited_fields": payload.edited_fields,
        "ai_adjusted_fields": payload.ai_adjusted_fields, "baseline_values": baseline,
        "baseline_source": baseline_source, "changed_fields": changed_fields,
        "locked_fields": locked_fields, "invalid_changed_fields": invalid_changed_fields,
        "allowed_fields": allowed, "instruction": payload.instruction or "", "values": values,
    }, ensure_ascii=False)
    if payload.operation == "regenerate_field":
        step = next(step for step in SETUP_FLOW_STEPS if step["key"] == payload.target_field)
        data, usage = await generate_validated_setup_options(
            db=db, project=project, current_user=current_user,
            step_key=payload.target_field, question=step["question"], values=values,
            context=json.dumps({"logline": project.logline or "", "user_instruction": payload.instruction or ""}, ensure_ascii=False),
            revision=payload.context_revision, template_instructions=template_instructions,
            action="regenerate_quick_setup_field",
        )
        return {"status": "options", "operation": payload.operation,
                "target_field": payload.target_field, **data, "tokens_used": usage,
                "total_tokens": int(project.total_tokens or 0), "context_revision": payload.context_revision}
    usage = 0
    raw_content = ""
    revised: Dict[str, str] = {}
    summary = ""
    try:
        revised, summary, usage = await invoke_with_quota(billed_user_id, lambda: llm.revise_quick_setup_fields(
            logline=project.logline or "", values=values, allowed_fields=allowed,
            instruction=payload.instruction or "", operation=payload.operation,
            scope=payload.scope, template_instructions=template_instructions,
            baseline_values=baseline, changed_fields=changed_fields,
            locked_fields=locked_fields, invalid_changed_fields=invalid_changed_fields,
        ))
        raw_content = getattr(revised, "raw_content", json.dumps({"fields": revised, "summary": summary}, ensure_ascii=False))
        setup_fields.validate_safety(revised, allowed=set(allowed))
        if not isinstance(summary, str) or len(summary) > MAX_INTERACTION_ANSWER_LENGTH:
            raise ValueError("AI 修订摘要过长或不是文本")
        normalized_revised = {field: setup_fields.normalize_field(field, value) for field, value in revised.items()}
        setup_fields.validate_safety({**values, **normalized_revised})
    except Exception as exc:
        if isinstance(exc, HTTPException) and exc.status_code == 429:
            raise
        usage = max(int(usage or 0), int(getattr(exc, "usage", 0) or 0))
        await record_quick_setup_ai_revision(
            db=db, project=project, current_user=current_user, operation=payload.operation,
            prompt=audit_prompt, response=str(getattr(exc, "raw_content", "") or raw_content),
            tokens=usage, status="failed",
            billed_user_id=billed_user_id,
            error_type=getattr(exc, "error_type", "invalid_ai_candidate"), error_message=str(exc),
        )
        raise HTTPException(status_code=503, detail="AI 修订结果无效，原草案未改变，请重试。") from exc
    stale = payload.context_revision != build_setup_context_revision(project)
    if not stale:
        try:
            await assert_setup_writable(db, project, payload.context_revision)
        except HTTPException as exc:
            if exc.status_code != 409:
                raise
            stale = True
    await record_quick_setup_ai_revision(
        db=db, project=project, current_user=current_user, operation=payload.operation,
        prompt=audit_prompt, response=raw_content, tokens=usage,
        status="stale" if stale else "success",
        billed_user_id=billed_user_id,
        error_type="stale_context" if stale else None,
        error_message="AI 返回前项目设定已更新或开始生成" if stale else None,
    )
    if stale:
        raise HTTPException(status_code=409, detail="AI 分析期间项目设定已更新，请刷新后重试。")
    if invalid_changed_fields:
        note = "锁定改后字段仍需人工修正，未自动改写：" + "、".join(invalid_changed_fields)
        summary = summary[:MAX_INTERACTION_ANSWER_LENGTH - len(note) - 1] + "\n" + note
    changes = [{"field": field, "before": values.get(field, ""), "after": value}
               for field, value in normalized_revised.items() if value != values.get(field, "")]
    return {"status": "candidate", "operation": payload.operation, "scope": payload.scope,
            "changes": changes, "changed_fields": [change["field"] for change in changes],
            "summary": summary.strip(), "tokens_used": max(0, int(usage or 0)),
            "total_tokens": int(project.total_tokens or 0), "context_revision": payload.context_revision}


@app.post("/projects/{project_id}/setup/quick-review")
async def submit_quick_setup_review(
    project_id: int, payload: QuickSetupReviewRequest,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user),
):
    project, _ = await require_project_access(db, project_id, current_user.id, minimum_role="editor")
    await assert_setup_writable(db, project, payload.context_revision)
    current_context = get_internal_project_context(project)
    status = payload.action
    if payload.action in {"save", "save_guided", "confirm"}:
        setup_drafts.require_current_draft(project)
    if payload.action in {"save", "save_guided"}:
        if "values" not in payload.model_fields_set:
            raise HTTPException(status_code=422, detail="保存工作稿需要明确提交 values；未提交时不会替换现有工作稿。")
        baseline, _ = setup_drafts.baseline_for(project, payload.baseline_values)
        previous, _ = setup_drafts.inspect_draft(project)
        ai_fields = payload.ai_adjusted_fields if "ai_adjusted_fields" in payload.model_fields_set else (previous or {}).get("ai_adjusted_fields", [])
        draft = setup_drafts.build_working_draft(
            project, payload.values, baseline, payload.edited_fields, ai_fields,
            mode_change=payload.action == "save_guided",
        )
        if payload.action == "save_guided":
            current_context[SETUP_MODE_KEY] = SETUP_MODE_GUIDED
            await write_setup(db, project, payload.context_revision, {"global_context": current_context, "quick_setup_draft": draft})
            status = "saved_guided"
        else:
            await write_working_draft(db, project, payload.context_revision, draft)
            status = "saved"
    elif payload.action == "discard":
        await write_working_draft(db, project, payload.context_revision, None)
        status = "discarded"
    elif payload.action in {"guided", "regenerate"}:
        current_context[SETUP_MODE_KEY] = SETUP_MODE_GUIDED if payload.action == "guided" else SETUP_MODE_AI_FAST
        if payload.action == "regenerate":
            current_context[setup_drafts.REGENERATE_KEY] = True
        else:
            current_context.pop(QUICK_EDITED_FIELDS_KEY, None)
            current_context.pop(setup_drafts.REGENERATE_KEY, None)
        await write_setup(db, project, payload.context_revision, {"global_context": current_context, "quick_setup_draft": None})
    else:
        try:
            normalized_values = normalize_quick_setup_values(project, payload.values, preserve_existing=False)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        for key in ("synopsis_brief", "synopsis_detailed", "brief_synopsis", "detailed_synopsis", "story_brief", "story_detailed", setup_drafts.REGENERATE_KEY):
            current_context.pop(key, None)
        baseline, _ = setup_drafts.baseline_for(project, payload.baseline_values)
        actual_changes = setup_drafts.value_changes(payload.values, baseline)
        current_context.update(normalized_values)
        current_context[SETUP_MODE_KEY] = SETUP_MODE_AI_FAST
        current_context[QUICK_EDITED_FIELDS_KEY] = [key for key in payload.edited_fields if key in actual_changes]
        current_context["final_confirm"] = "confirmed"
        await write_setup(db, project, payload.context_revision, {
            "global_context": current_context, "project_type": normalized_values["project_type"],
            "title": normalized_values["title"], "quick_setup_draft": None,
        })
        status = "confirmed"
    await db.commit()
    return {
        "status": status, "setup_mode": get_internal_project_context(project).get(SETUP_MODE_KEY, ""),
        "context": project.global_context, "title": project.title or "",
        "total_tokens": int(project.total_tokens or 0), **revision_meta(project), **draft_response_state(project),
    }


@app.post("/projects/{project_id}/analyze")
async def analyze_logline(
    project_id: int, 
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    """
    Phase 1: Deep Analysis & Setup.
    Iteratively helps the user build the 'Project Bible' by asking questions.
    """
    project, _ = await require_project_access(
        db,
        project_id,
        current_user.id,
        minimum_role="editor",
    )
    await db.refresh(project)

    def with_runtime_meta(payload: Dict[str, Any]) -> Dict[str, Any]:
        result = dict(payload or {})
        result.pop("_setup_cache", None)
        result.update(revision_meta(project))
        if isinstance(result.get("payload"), dict):
            result["payload"] = {**result["payload"], "context_revision": build_setup_context_revision(project)}
            if result["payload"].get("field") == QUICK_REVIEW_FIELD:
                data = result["payload"]
                values = data.setdefault("values", {item["key"]: item["value"] for item in data.get("sections", [])})
                data.setdefault("baseline_values", dict(values))
                data.setdefault("edited_fields", [])
                data.setdefault("ai_adjusted_fields", [])
                data.setdefault("draft_status", "generated")
                data.setdefault("draft_stale", False)
                data.setdefault("read_only", False)
                data.setdefault("base_setup_revision", int(project.setup_revision or 0))
                data.setdefault("saved_at", None)
        result["total_tokens"] = int(project.total_tokens or 0)
        result["setup_mode"] = str(
            get_internal_project_context(project).get(SETUP_MODE_KEY, "") or ""
        )
        saved, stale = setup_drafts.inspect_draft(project)
        result["saved_draft_available"] = saved is not None
        result["draft_stale"] = stale
        return result

    saved_draft, saved_stale = setup_drafts.inspect_draft(project)
    saved_mode = get_internal_project_context(project).get(SETUP_MODE_KEY, "")
    if saved_draft is not None and saved_mode != SETUP_MODE_GUIDED:
        return with_runtime_meta(working_draft_payload(project, saved_draft, stale=saved_stale))

    await assert_setup_writable(db, project, build_setup_context_revision(project))
    billed_user_id = project.owner_id
    normalization_revision = build_setup_context_revision(project)
    normalized_draft = detached_setup(project)
    normalized = False
    if saved_draft is None and normalize_project_title(normalized_draft):
        normalized = True
    if saved_draft is None and normalize_project_context(normalized_draft):
        normalized = True
    if normalized:
        await write_setup(db, project, normalization_revision, {
            "title": normalized_draft.title, "global_context": normalized_draft.global_context,
        })
        await db.commit()

    normalized_context = build_normalized_context(project)
    internal_context = get_internal_project_context(project)
    setup_mode = str(internal_context.get(SETUP_MODE_KEY, "") or "").strip()
    if setup_mode not in {SETUP_MODE_AI_FAST, SETUP_MODE_GUIDED}:
        if has_existing_setup_progress(project):
            internal_context[SETUP_MODE_KEY] = SETUP_MODE_GUIDED
            await write_setup(db, project, build_setup_context_revision(project), {"global_context": internal_context})
            setup_mode = SETUP_MODE_GUIDED
            await db.commit()
        else:
            return with_runtime_meta(
                {
                    "type": "interaction_required",
                    "payload": {
                        "field": SETUP_MODE_FIELD,
                        "question": "你希望如何完善故事设定？",
                        "options": [
                            {
                                "label": "✨ AI 快速完成",
                                "value": SETUP_MODE_AI_FAST,
                                "description": "AI 生成一份完整且连贯的设定草案，你只需检查和修改有疑问的内容。",
                            },
                            {
                                "label": "🎛️ 自己掌控",
                                "value": SETUP_MODE_GUIDED,
                                "description": "逐项决定剧情、人物、主题和视觉方向，完整控制每一个创作决定。",
                            },
                        ],
                    },
                }
            )

    if saved_draft is None and should_auto_prefill_from_logline(project, normalized_context):
        await enforce_user_quota(db, billed_user_id)
        prefill_revision = build_setup_context_revision(project)
        prefill_draft = detached_setup(project)
        filled_fields: List[str] = []
        prefill_changed = False
        prefill_usage = 0
        extracted_setup: Dict[str, Any] = {}
        prefill_error = None
        prefill_raw = ""
        try:
            extracted_setup, prefill_usage = await invoke_with_quota(billed_user_id, lambda: llm.extract_setup_from_long_input(project.logline or ""))
            prefill_raw = getattr(extracted_setup, "raw_content", json.dumps(extracted_setup, ensure_ascii=False))
            filled_fields, prefill_changed = apply_auto_prefill(prefill_draft, extracted_setup)
        except Exception as exc:
            if isinstance(exc, HTTPException) and exc.status_code == 429:
                raise
            prefill_error = exc
            prefill_usage = max(int(prefill_usage or 0), int(getattr(exc, "usage", 0) or 0))
            prefill_raw = str(getattr(exc, "raw_content", "") or prefill_raw)
            logger.warning(f"Failed to auto-prefill setup for project {project_id}: {exc}")
            filled_fields, prefill_changed = apply_auto_prefill(prefill_draft, {})
        if prefill_usage:
            await increment_project_tokens(db, project, prefill_usage)
            await db.commit()
        stale = False
        try:
            await assert_setup_writable(db, project, prefill_revision)
        except HTTPException as exc:
            if exc.status_code != 409:
                raise
            stale = True
        await write_setup_ai_audit(
            user_id=current_user.id, project_id=project_id, action="auto_prefill_setup",
            billed_user_id=billed_user_id,
            prompt=project.logline or "", response=prefill_raw, tokens=prefill_usage,
            status="stale" if stale else "failed" if prefill_error else "success",
            error_type="stale_context" if stale else getattr(prefill_error, "error_type", type(prefill_error).__name__) if prefill_error else None,
            error_message="预填返回前项目设定已变化" if stale else str(prefill_error) if prefill_error else None,
        )
        if stale:
            raise HTTPException(status_code=409, detail="AI 预填期间项目设定已更新，请刷新后重试。")

        if filled_fields:
            logger.info(f"项目 {project_id} 已从长输入自动补全字段: {', '.join(filled_fields)}")
        if prefill_changed:
            await write_setup(db, project, prefill_revision, {
                "global_context": prefill_draft.global_context,
                "title": prefill_draft.title,
                "project_type": prefill_draft.project_type,
            })
            await db.commit()

        normalized_context = build_normalized_context(project)

    # A legacy/mismatched cache is ignored, never relabelled as current.
    quick_review_needed = setup_mode == SETUP_MODE_AI_FAST and ("final_confirm" not in normalized_context or bool(get_internal_project_context(project).get(setup_drafts.REGENERATE_KEY)))
    stage = QUICK_REVIEW_FIELD if quick_review_needed else next(
        (step["key"] for step in get_relevant_setup_steps(normalized_context.get("project_type", "movie"))
         if step["key"] not in normalized_context), "completed"
    )
    if has_valid_setup_cache(project, mode=setup_mode, stage=stage):
        logger.info(f"项目 {project_id} 命中缓存，直接返回之前的提问。")
        return with_runtime_meta(project.next_step_cache)
    analysis_revision = build_setup_context_revision(project)

    logger.info(f"正在分析项目 {project_id} 的进度状况...")

    if quick_review_needed:
        field_specs = quick_setup_field_specs()
        interaction_template = await get_prompt_addendum(
            db,
            stage="interaction",
            project_type=(
                project.project_type
                if project.project_type and project.project_type != "pending"
                else "all"
            ),
        )
        draft_values: Optional[Dict[str, str]] = None
        draft_usage = 0
        last_error: Optional[Exception] = None
        for attempt in range(1, MAX_INTERACTION_ATTEMPTS + 1):
            await enforce_user_quota(db, billed_user_id)
            attempt_usage = 0
            raw_content = ""
            error_type = None
            try:
                generated_values, attempt_usage = await invoke_with_quota(billed_user_id, lambda: llm.generate_quick_setup_draft(
                    logline=project.logline or "", current_context=normalized_context,
                    field_specs=field_specs, template_instructions=interaction_template,
                ))
                raw_content = getattr(generated_values, "raw_content", json.dumps(generated_values, ensure_ascii=False))
                draft_values = normalize_quick_setup_values(project, generated_values)
            except Exception as exc:
                if isinstance(exc, HTTPException) and exc.status_code == 429:
                    raise
                last_error = exc
                draft_values = None
                attempt_usage = max(int(attempt_usage or 0), int(getattr(exc, "usage", 0) or 0))
                raw_content = str(getattr(exc, "raw_content", "") or raw_content)
                error_type = getattr(exc, "error_type", type(exc).__name__)
            attempt_usage = max(0, int(attempt_usage or 0))
            draft_usage += attempt_usage
            if attempt_usage:
                await increment_project_tokens(db, project, attempt_usage)
                await db.commit()
            stale = False
            try:
                await assert_setup_writable(db, project, analysis_revision)
            except HTTPException as exc:
                if exc.status_code != 409:
                    raise
                stale = True
            await write_setup_ai_audit(
                user_id=current_user.id, project_id=project_id,
                billed_user_id=billed_user_id,
                action="generate_quick_setup_draft", prompt=project.logline or "",
                response=raw_content, tokens=attempt_usage,
                status="stale" if stale else "success" if draft_values is not None else "failed",
                step_key=QUICK_REVIEW_FIELD, attempt=attempt,
                error_type="stale_context" if stale else error_type,
                error_message="AI 返回前项目设定已更新" if stale else str(last_error) if draft_values is None else None,
            )
            if stale:
                raise HTTPException(status_code=409, detail="AI 分析期间项目设定已更新，请刷新后重试。")
            if draft_values is not None:
                break

        if draft_values is None:
            await assert_setup_writable(db, project, analysis_revision)
            return with_runtime_meta(
                {
                    "type": "interaction_required",
                    "payload": {
                        "field": SETUP_MODE_FIELD,
                        "question": (
                            "AI 快速设定暂时生成失败，你可以重试或切换到自己掌控。"
                            if last_error is not None
                            else "AI 未返回有效草案，请选择下一步。"
                        ),
                        "options": [
                            {
                                "label": "✨ 重试 AI 快速完成",
                                "value": SETUP_MODE_AI_FAST,
                                "description": "重新联合生成一份完整故事设定草案。",
                            },
                            {
                                "label": "🎛️ 切换到自己掌控",
                                "value": SETUP_MODE_GUIDED,
                                "description": "保留已经确认的内容，从下一项缺失设定继续。",
                            },
                        ],
                    },
                }
            )

        response_payload = {
            "type": "interaction_required",
            "payload": {
                "field": QUICK_REVIEW_FIELD,
                "question": "AI 已完成整套故事设定，请展开有疑问的内容进行修改。",
                "context_revision": build_setup_context_revision(project),
                "sections": build_quick_review_sections(project, draft_values),
                "values": dict(draft_values), "baseline_values": dict(draft_values),
                "edited_fields": [], "ai_adjusted_fields": [],
                "draft_status": "generated", "draft_stale": False, "read_only": False,
                "base_setup_revision": int(project.setup_revision or 0), "saved_at": None,
            },
        }
        response_payload = await write_setup_cache(
            db, project, analysis_revision, response_payload,
            mode=setup_mode, stage=QUICK_REVIEW_FIELD,
        )
        await db.commit()
        return with_runtime_meta(response_payload)

    # 1. Check which steps are missing
    # Important: 'project_type' is stored in column, others in global_context
    normalized_context = build_normalized_context(project)
    
    # Calculate Total Steps (Dynamic based on Type)
    p_type = normalized_context.get("project_type", "movie")
    relevant_steps = get_relevant_setup_steps(p_type)

    next_step = None
    next_step_index = 0
    total_steps = len(relevant_steps)

    for i, step in enumerate(relevant_steps):
        if step["key"] not in normalized_context:
            next_step = step
            next_step_index = i + 1
            break
            
    # 2. If all steps completed -> Proceed to Outline Generation
    if not next_step:
        logger.info(f"项目 {project_id} 所有基础设定步骤已完成，准备生成大纲。")
        return with_runtime_meta({"type": "completed", "message": "基础设定已完成！准备生成大纲..."})

    logger.info(f"项目 {project_id} 下一步骤: {next_step['key']} ({next_step_index}/{total_steps})")
    
    # helper to inject progress info
    def add_progress(payload):
        payload["progress"] = {"current": next_step_index, "total": total_steps}
        return payload

    # 3. Handle specific logic for the next step
    # 3.1 Hardcoded options for Type
    if next_step["key"] == "project_type":
        return with_runtime_meta({
            "type": "interaction_required",
            "payload": add_progress({
                "field": "project_type",
                "question": next_step["question"],
                "options": next_step["default_options"]
            })
        })
    
    # 3.2 Hardcoded options for Episode Count / Movie Duration / Scene Count
    if next_step["key"] == "movie_duration":
         return with_runtime_meta({
            "type": "interaction_required",
            "payload": add_progress({
                "field": "movie_duration",
                "question": next_step["question"],
                "options": [
                    {"label": "90分钟 (标准电影)", "value": "90"},
                    {"label": "120分钟 (长篇商业片)", "value": "120"},
                    {"label": "150分钟以上 (史诗篇幅)", "value": "150"},
                    {"label": "60分钟 (中片/电视电影)", "value": "60"}
                ]
            })
        })

    if next_step["key"] == "scene_count_target":
         return with_runtime_meta({
            "type": "interaction_required",
            "payload": add_progress({
                "field": "scene_count_target",
                "question": next_step["question"],
                "options": [
                    {"label": "40场 (简约大纲)", "value": "40"},
                    {"label": "60场 (标准大纲)", "value": "60"},
                    {"label": "100场 (精细大纲)", "value": "100"},
                    {"label": "120场以上 (极度详尽)", "value": "120"}
                ]
            })
        })

    if next_step["key"] == "episode_count":
         return with_runtime_meta({
            "type": "interaction_required",
            "payload": add_progress({
                "field": "episode_count",
                "question": next_step["question"],
                "options": [
                    {"label": "8集 (迷你剧)", "value": "8"},
                    {"label": "12集 (标准季)", "value": "12"},
                    {"label": "20集 (国产剧标准)", "value": "20"},
                    {"label": "24集", "value": "24"},
                    {"label": "40集以上", "value": "40"}
                ]
            })
        })
    
    if next_step["key"] == "episode_duration":
         return with_runtime_meta({
            "type": "interaction_required",
            "payload": add_progress({
                "field": "episode_duration",
                "question": next_step["question"],
                "options": [
                    {"label": "1-2分钟 (竖屏短剧)", "value": "2mins"},
                    {"label": "5-10分钟 (迷你剧)", "value": "10mins"},
                    {"label": "20分钟 (情景喜剧/动画)", "value": "20mins"},
                    {"label": "45分钟 (标准剧集)", "value": "45mins"},
                    {"label": "60分钟 (美剧/电影感)", "value": "60mins"}
                ]
            })
        })
    
    if next_step.get("is_confirmation"):
        normalized_context = await ensure_story_synopsis(
            project,
            normalized_context,
            db,
            persist=False,
            actor_id=current_user.id,
            optional=True,
        )
        # Synopsis enrichment is a setup write as well; preserve internal fields.
        enriched_context = {**get_internal_project_context(project), **normalized_context}
        if enriched_context != get_internal_project_context(project):
            await db.commit()  # preserve already-used AI tokens even if CAS rejects stale output
            await write_setup(db, project, analysis_revision, {"global_context": enriched_context})
        if project.project_type and project.project_type != "pending":
            normalized_context["project_type"] = project.project_type

        summary_text = build_context_summary(project, normalized_context)
        await db.commit()
        return with_runtime_meta({
            "type": "interaction_required",
            "payload": add_progress({
                "field": "final_confirm",
                "question": next_step["question"],
                "context_summary": summary_text,
                "options": [
                    {"label": "✅ 确定并开始生成", "value": "confirmed"},
                    *[
                        {"label": label, "value": f"edit:{target_key}"}
                        for target_key, label in FINAL_CONFIRM_EDIT_TARGETS
                    ],
                    {"label": "🔄 重新设定 (清空当前设定重头开始)", "value": "reset"}
                ]
            })
        })

    if next_step["key"] == "video_duration_seconds":
         return with_runtime_meta({
            "type": "interaction_required",
            "payload": add_progress({
                "field": "video_duration_seconds",
                "question": next_step["question"],
                "options": [
                    {"label": "15秒（1条提示词）", "value": "15"},
                    {"label": "30秒（2条提示词）", "value": "30"},
                    {"label": "45秒（3条提示词）", "value": "45"},
                    {"label": "60秒（4条提示词）", "value": "60"},
                    {"label": "90秒（6条提示词）", "value": "90"},
                    {"label": "120秒（8条提示词）", "value": "120"}
                ]
            })
        })
    
    # 3.4 Check Prompt Richness (Optimization)
    # If the user's initial logline is very long (> 100 chars) and detailed,
    # we tell the LLM to verify if we even need to ask this question.
    # Note: Currently we just proceed to ask to be comprehensive.
    
    # 3.4 For other steps, use LLM to generate context-aware options
    # We pass the logline + current context to LLM
    prompt_context = f"Logline: {project.logline}\nCurrent Settings: {json.dumps(normalized_context, ensure_ascii=False)}"
    
    logger.info(f"正在调用 LLM 为步骤 {next_step['key']} 生成选项...")
    interaction_template = await get_prompt_addendum(
        db,
        stage="interaction",
        project_type=project.project_type,
    )
    
    question_data, usage = await generate_validated_setup_options(
        db=db, project=project, current_user=current_user,
        step_key=next_step["key"], question=next_step["question"],
        values={key: str(value) for key, value in normalized_context.items() if key in setup_fields.SETUP_FIELDS},
        context=prompt_context, revision=analysis_revision,
        template_instructions=interaction_template,
        action=f"analyze_step_{next_step['key']}",
    )
    
    # Construction Response
    response_payload = {
        "type": "interaction_required",
        "payload": add_progress({
            "field": next_step["key"],
            "question": question_data.get("question", next_step["question"]), 
            "options": question_data.get("options", [])
        })
    }

    if next_step["key"] == "title":
        response_payload["payload"]["options"] = sanitize_title_options(
            response_payload["payload"].get("options", [])
        )
    
    # Cache the result to DB so next fetch is instant
    response_payload = await write_setup_cache(
        db, project, analysis_revision, response_payload,
        mode=setup_mode, stage=next_step["key"],
    )
    await db.commit()

    return with_runtime_meta(response_payload)

@app.post("/projects/{project_id}/generate_scenes")
async def generate_scenes(
    project_id: int,
    selected_option: Optional[str] = Query(default=None, max_length=1000),
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    """
    Phase 1.5: User selected an option, now generate outline.
    Phase 2: Persist a durable generation job for the worker.
    """
    logger.info(f"收到生成分场大纲请求，项目ID: {project_id}")
    # 1. Update project genre/style based on selected_option
    project = await db.get(models.Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    role = await project_role(db, project, current_user.id)
    if role not in {"owner", "editor"}:
        raise HTTPException(status_code=404, detail="Project not found")
    project.access_role = role
    if not await claim_generation(db, project_id, current_user.id):
        await db.rollback()
        raise HTTPException(status_code=409, detail="该项目已有生成任务正在运行")
    await db.commit()
    await db.refresh(project)

    try:
        c = await ensure_story_synopsis(project, project.global_context or {}, db, actor_id=current_user.id)
    except Exception:
        await mark_claimed_project_failed(db, project_id)
        raise
    project.global_context = c

    # Use selected_option if string generic, or fallback to stored context values
    style_context = str(selected_option or "").strip()
    if not style_context or style_context.lower() == "auto":
        # Construct summary from context
        style_context = f"Genre: {project.project_type}, Tone: {c.get('tone')}, Style: {c.get('visual_style')}"

    # Canonical durations retain exact decimal/unit semantics downstream.
    target_count = 5
    duration_seconds = 0
    count_key = "scene_count_target" if project.project_type == "movie" else "episode_count"
    try:
        if project.project_type == "short_video":
            duration_seconds = int(setup_fields.normalize_number("video_duration_seconds", str(c.get("video_duration_seconds") or "60")))
            target_count = (duration_seconds + 14) // 15
        elif c.get(count_key):
            target_count = int(setup_fields.normalize_number(count_key, str(c[count_key])))
        elif project.project_type == "movie" and c.get("movie_duration"):
            minutes = Decimal(setup_fields.normalize_number("movie_duration", str(c["movie_duration"])))
            # A count estimate deliberately floors the exact 0.8 scenes/minute;
            # the stored/displayed duration is never rounded or truncated.
            target_count = int((minutes * Decimal("0.8")).to_integral_value(rounding=ROUND_FLOOR))
    except ValueError as exc:
        await mark_claimed_project_failed(db, project_id)
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    target_limit = GENERATION_TARGET_LIMITS.get(project.project_type or "", 100)
    if target_count > target_limit:
        logger.warning(
            "项目 %s 请求生成 %s 个条目，已限制为 %s",
            project_id,
            target_count,
            target_limit,
        )
    target_count = max(1, min(int(target_count or 1), target_limit))

    project.genre = style_context
    clear_generation_error(project)
    invalidate_scene_prompt_cache(project)

    if project.project_type == "short_video":
        style_context = (
            f"{style_context}; 模式:短视频15秒分段提示词; 总时长:{duration_seconds}秒;"
            f" 需要生成{target_count}条15秒提示词"
        )

    # Force clearing of any old scenes from a previous attempt
    try:
        existing_scene_count = int(
            await db.scalar(
                select(func.count())
                .select_from(models.Scene)
                .where(models.Scene.project_id == project_id)
            )
            or 0
        )
        if existing_scene_count:
            await create_project_version(
                db,
                project_id,
                current_user.id,
                "重新生成前自动快照",
            )
        await db.execute(delete(models.Scene).where(models.Scene.project_id == project_id))
        job = await enqueue_job(
            db,
            project_id=project_id,
            kind=OUTLINE_JOB,
            payload={
                "style_context": style_context,
                "target_count": target_count,
                "user_id": current_user.id,
            },
        )
        await db.commit()
    except Exception as exc:
        await mark_claimed_project_failed(db, project_id)
        logger.exception("Failed to prepare project %s for generation: %s", project_id, exc)
        raise HTTPException(status_code=500, detail="生成任务准备失败，请稍后重试")

    logger.info(
        "生成大纲任务已入队。job=%s project=%s count=%s",
        job.id,
        project_id,
        target_count,
    )

    return {
        "status": "Scene generation queued",
        "project_id": project_id,
        "job_id": job.id,
    }

# --- Background Task Implementation ---

async def run_incremental_outline_generation(project_id: int, style_context: str, target_count: int, user_id: int):
    logger.info(f"[Task] Starting Incremental Outline Gen for Project {project_id}")
    
    async with database.SessionLocal() as db:
        project = await db.get(models.Project, project_id)
        if not project:
            return
        
        # Determine Batch Size (User requested "safe/one-by-one", so we choose 1 to be absolutely safe and responsive)
        # Using 1 allows frontend to see each scene pop up.
        batch_size = 1 
        current_idx = 1
        generated_scenes: list[models.Scene] = []
        story_bible = build_story_bible(
            logline=project.logline,
            project_type=project.project_type,
            genre=style_context,
            global_context=project.global_context or {},
        )
        template_instructions = await get_prompt_addendum(
            db,
            stage="outline",
            project_type=project.project_type,
        )
        
        try:
            while current_idx <= target_count:
                # Re-check status in case user cancelled or a restart recovered the job.
                await db.refresh(project)
                if project.status != models.ProcessingStatus.GENERATING:
                    logger.info("[Task] Outline generation no longer active.")
                    return

                end_idx = min(current_idx + batch_size - 1, target_count)
                logger.info(f"[Task] Generating scenes {current_idx}-{end_idx}...")
                continuity_context = build_outline_continuity_context(
                    story_bible=story_bible,
                    prior_scenes=generated_scenes,
                    current_index=current_idx,
                    total_scenes=target_count,
                )
                batch_scenes, usage = await run_project_ai_call(
                    db=db, project=project, actor_id=user_id, action=f"outline_scene_{current_idx}",
                    prompt=continuity_context, expected_status=models.ProcessingStatus.GENERATING,
                    validate=validate_outline_batch,
                    result_status=lambda value: "partial" if looks_like_story_restart(value[0]["outline"], current_idx) else "success",
                    invoke=lambda: llm.generate_scene_batch(
                        project.logline,
                        style_context,
                        current_idx,
                        end_idx,
                        previous_context=continuity_context,
                        total_target=target_count,
                        template_instructions=template_instructions,
                    ),
                )
                await db.refresh(project)
                if project.status != models.ProcessingStatus.GENERATING:
                    logger.info(
                        "[Task] Outline result for scene %s discarded after cancellation.",
                        current_idx,
                    )
                    return

                first_outline = str(
                    (batch_scenes or [{}])[0].get("outline", "") or ""
                )
                if looks_like_story_restart(first_outline, current_idx):
                    logger.warning(
                        "[Task] Scene %s looked like a story restart; regenerating once.",
                        current_idx,
                    )
                    guarded_context = build_outline_continuity_context(
                        story_bible=story_bible,
                        prior_scenes=generated_scenes,
                        current_index=current_idx,
                        total_scenes=target_count,
                        extra_warning=(
                            "上一次结果疑似重新开篇。必须改为承接上一场的新事件，"
                            "禁止使用‘故事开始、序幕、初次相遇、第一次见面’等重启表达。"
                        ),
                    )
                    retry_scenes, retry_usage = await run_project_ai_call(
                        db=db, project=project, actor_id=user_id, action=f"outline_scene_{current_idx}",
                        prompt=guarded_context, attempt=2, expected_status=models.ProcessingStatus.GENERATING,
                        validate=lambda value: validate_continuation(value, current_idx, outline=True),
                        invoke=lambda: llm.generate_scene_batch(
                            project.logline,
                            style_context,
                            current_idx,
                            end_idx,
                            previous_context=guarded_context,
                            total_target=target_count,
                            template_instructions=template_instructions,
                        ),
                    )
                    await db.refresh(project)
                    if project.status != models.ProcessingStatus.GENERATING:
                        logger.info(
                            "[Task] Retried outline for scene %s discarded after cancellation.",
                            current_idx,
                        )
                        return
                    usage += int(retry_usage or 0)
                    batch_scenes = retry_scenes
                    first_outline = str(
                        (batch_scenes or [{}])[0].get("outline", "") or ""
                    )
                    if looks_like_story_restart(first_outline, current_idx):
                        raise RuntimeError(
                            "连续性守卫拒绝了疑似重新开篇的场次大纲"
                        )

                if len(batch_scenes or []) != batch_size:
                    raise RuntimeError(
                        f"Outline batch {current_idx}-{end_idx} returned "
                        f"{len(batch_scenes or [])} items; expected {batch_size}."
                    )

                # Enforce strictly sequential indexing; never trust the model's index.
                for offset, scene_data in enumerate(batch_scenes):
                    outline = str(scene_data.get("outline", "") or "").strip()
                    if not outline:
                        raise RuntimeError(
                            f"Outline batch {current_idx}-{end_idx} returned empty content."
                        )
                    new_scene = models.Scene(
                        project_id=project.id,
                        scene_index=current_idx + offset,
                        outline=outline,
                        status=models.ProcessingStatus.PENDING
                    )
                    db.add(new_scene)
                    generated_scenes.append(new_scene)

                await db.commit()
                current_idx += batch_size
        except Exception as exc:
            await db.rollback()
            if isinstance(exc, HTTPException) and exc.status_code == 409:
                return
            logger.exception(f"[Task] Outline generation failed: {exc}")
            project = await db.get(models.Project, project_id)
            if project:
                project.status = models.ProcessingStatus.FAILED
                record_generation_error(project, exc, stage="outline")
                await db.commit()
            return

        # After Outline Complete -> Trigger Content Generation
        logger.info("[Task] Outline Complete. Starting Content Gen Loop...")
        await run_generation_loop(project.id, user_id=user_id)


@app.post("/projects/{project_id}/scenes/{scene_index}/regenerate")
async def regenerate_scene(
    project_id: int, 
    scene_index: int,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    project, _ = await require_project_access(
        db, project_id, current_user.id, minimum_role="editor"
    )
    result = await db.execute(
        select(models.Scene)
        .where(models.Scene.project_id == project_id)
        .where(models.Scene.scene_index == scene_index)
    )
    scene = result.scalars().first()
    if not scene:
        raise HTTPException(status_code=404, detail="Scene not found")

    if not await claim_generation(db, project_id, current_user.id):
        await db.rollback()
        raise HTTPException(status_code=409, detail="该项目已有生成任务正在运行")

    # The project may have changed between access lookup and the atomic claim.
    # Refresh under the claim's write lock before clearing derived context.
    await db.refresh(project)
    # A restore can delete/recreate scenes (including reusing an old row id).
    # Resolve the requested logical scene again while holding the claim lock.
    scene = await db.scalar(
        select(models.Scene)
        .where(models.Scene.project_id == project_id)
        .where(models.Scene.scene_index == scene_index)
        .execution_options(populate_existing=True)
    )
    if scene is None:
        await db.rollback()
        raise HTTPException(status_code=404, detail="场次已变化或不存在，请刷新后重试")
    await create_project_version(
        db,
        project_id,
        current_user.id,
        f"重写第{scene_index}场前自动快照",
    )

    # Reset status
    scene.status = models.ProcessingStatus.PENDING
    scene.content = None
    scene.summary = None
    project.status = models.ProcessingStatus.GENERATING
    invalidate_scene_prompt_cache(project, scene_index)

    try:
        job = await enqueue_job(
            db,
            project_id=project_id,
            kind=CONTENT_JOB,
            payload={"scene_index": scene_index, "user_id": current_user.id},
        )
        await db.commit()
    except Exception as exc:
        await mark_claimed_project_failed(db, project_id)
        logger.exception(
            "Failed to enqueue scene regeneration for project %s: %s",
            project_id,
            exc,
        )
        raise HTTPException(status_code=500, detail="重新生成任务准备失败，请稍后重试")

    return {
        "status": "Regeneration queued",
        "project_id": project_id,
        "scene_index": scene_index,
        "job_id": job.id,
    }

@app.post("/projects/{project_id}/scenes/{scene_index}/to_prompt")
async def rewrite_scene_to_prompt(
    project_id: int,
    scene_index: int,
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    project, _ = await require_project_access(
        db, project_id, current_user.id, minimum_role="editor"
    )
    result = await db.execute(
        select(models.Scene)
        .where(models.Scene.project_id == project_id)
        .where(models.Scene.scene_index == scene_index)
    )
    scene = result.scalars().first()
    if not scene:
        raise HTTPException(status_code=404, detail="Scene not found")
    if (
        scene.status != models.ProcessingStatus.COMPLETED
        or not str(scene.content or "").strip()
    ):
        raise HTTPException(status_code=409, detail="场次内容尚未生成完成")

    context = dict(project.global_context) if isinstance(project.global_context, dict) else {}
    prompt_cache = context.get("_scene_ai_prompts")
    if not isinstance(prompt_cache, dict):
        prompt_cache = {}

    cache_key = str(scene_index)
    cached_prompt = str(prompt_cache.get(cache_key, "") or "").strip()
    if cached_prompt:
        return {"scene_index": scene_index, "prompt": cached_prompt, "cached": True}

    prompt_revision = build_setup_context_revision(project)
    prompt_scene_id = scene.id
    prompt_scene_outline = scene.outline
    prompt_scene_content = scene.content
    prompt_actor_id = current_user.id
    rewrite_prompt = (
        f"project_id={project_id}, scene_index={scene_index}, "
        f"outline={scene.outline or ''}, content={scene.content or ''}"
    )
    prompt_template = await get_prompt_addendum(
        db,
        stage="prompt",
        project_type=project.project_type,
    )

    async def prompt_is_stale():
        current_scene = await db.scalar(select(models.Scene).where(models.Scene.id == prompt_scene_id).execution_options(populate_existing=True))
        return (build_setup_context_revision(project) != prompt_revision
                or project.status == models.ProcessingStatus.GENERATING
                or current_scene is None or current_scene.status != models.ProcessingStatus.COMPLETED
                or current_scene.outline != prompt_scene_outline or current_scene.content != prompt_scene_content)

    def validate_prompt(value):
        if not str(value or "").strip():
            raise ValueError("AI 未返回有效提示词")

    prompt_audit_ids: List[int] = []
    try:
        prompt_text, usage = await run_project_ai_call(
            db=db, project=project, actor_id=prompt_actor_id,
            action=f"scene_to_prompt_{scene_index}", prompt=rewrite_prompt,
            stale_check=prompt_is_stale, validate=validate_prompt,
            audit_ids=prompt_audit_ids,
            invoke=lambda: llm.rewrite_scene_to_ai_prompt(
                project_type=project.project_type or "movie",
                logline=project.logline or "",
                style_guide=project.genre or "",
                scene_outline=scene.outline or "",
                scene_content=scene.content or "",
                scene_index=scene_index,
                template_instructions=prompt_template,
            ),
        )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=503, detail="AI 提示词转写失败，请稍后重试") from exc

    final_prompt = str(prompt_text or "").strip()
    try:
        await write_scene_prompt_cache(
            db, project, prompt_revision,
            scene_id=prompt_scene_id, scene_index=scene_index,
            scene_outline=prompt_scene_outline, scene_content=prompt_scene_content,
            prompt=final_prompt,
        )
        await db.commit()
    except HTTPException as exc:
        if exc.status_code != 409:
            raise
        # CAS is the final arbiter: if it lost a race after pre-audit validation,
        # revise the existing row, never append another charged usage record.
        await db.rollback()
        try:
            for audit_id in prompt_audit_ids:
                await update_ai_action_status(audit_id, status="stale", error_type="stale_context",
                                              error_message="提示词返回后、保存前项目设定或场次内容已变化")
        except Exception as audit_exc:
            raise HTTPException(status_code=503, detail="AI 审计最终状态写入失败，本次结果未应用，请管理员检查。") from audit_exc
        raise HTTPException(status_code=409, detail="提示词生成期间项目已更新，请刷新后重试。") from exc

    return {"scene_index": scene_index, "prompt": final_prompt, "cached": False}

# --- Export ---
@app.get("/projects/{project_id}/export")
async def export_project(
    project_id: int, 
    format: str = "txt",
    db: AsyncSession = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    # Eager load scenes
    result = await db.execute(
        select(models.Project)
        .where(models.Project.id == project_id)
        .options(selectinload(models.Project.scenes))
    )
    project = result.scalars().first()
    
    if not project or not await project_role(db, project, current_user.id):
        raise HTTPException(status_code=404, detail="Project not found")

    filename_raw = project.title or 'Untitled_Script'
    filename_encoded = quote(filename_raw)
    
    # Prepare Content Data
    project_scenes = sorted(project.scenes, key=lambda s: s.scene_index)
    context = project.global_context or {}
    
    if format == "docx":
        if not DocxDocument:
            raise HTTPException(501, "Word export library (python-docx) not installed on server.")
        
        doc = DocxDocument()
        doc.add_heading(project.title or "Untitled", 0)
        
        doc.add_heading("Project Bible", level=1)
        doc.add_paragraph(f"Logline: {project.logline}")
        doc.add_paragraph(f"Type: {project.project_type} | Genre: {project.genre}")
        for k, v in context.items():
            if k not in ['logline', 'project_type']:
                try:
                     doc.add_paragraph(f"{str(k).capitalize()}: {str(v)}")
                except:
                     pass
                
        doc.add_page_break()
        doc.add_heading("Screenplay", level=1)
        
        for scene in project_scenes:
            doc.add_heading(f"SCENE {scene.scene_index}", level=2)
            doc.add_paragraph(f"Outline: {scene.outline}", style='Intense Quote')
            if scene.content:
                # Basic formatting for script
                doc.add_paragraph(scene.content)
            else:
                doc.add_paragraph("[Content Generating...]")
            doc.add_paragraph("") # Spacing
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename_encoded}.docx"}
        )

    elif format == "md":
        content = f"# {project.title or 'Untitled'}\n\n"
        content += f"**Logline:** {project.logline}\n\n"
        content += f"**Type:** {project.project_type}\n"
        content += "---\n\n## Project Settings\n"
        for k, v in context.items():
             content += f"- **{k}:** {v}\n"
        content += "\n---\n\n## Script\n\n"
        
        for scene in project_scenes:
            content += f"### SCENE {scene.scene_index}\n"
            content += f"> **Outline:** {scene.outline}\n\n"
            content += (scene.content or "[Generating...]") + "\n\n"
            content += "---\n\n"
            
        return StreamingResponse(
            io.BytesIO(content.encode('utf-8')),
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename_encoded}.md"}
        )
        
    else: # Default TXT
        content = f"Title: {project.title}\nLogline: {project.logline}\n\n"
        for scene in project_scenes:
            content += f"SCENE {scene.scene_index}\n{scene.content or ''}\n\n"
        return StreamingResponse(
            io.BytesIO(content.encode('utf-8')),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename_encoded}.txt"}
        )

# --- Background Task (The Engine) ---

async def _run_generation_loop(project_id: int, user_id: Optional[int] = None):
    """
    The Core Loop: Iterates scenes and generates content with Rolling Summary.
    """
    logger.info(f"[后台任务] 开始为项目 {project_id} 生成剧本内容...")
    
    async with database.SessionLocal() as db:
        project = await db.get(models.Project, project_id)
        if not project: 
            logger.error(f"[后台任务] 项目 {project_id} 未找到，任务中止")
            return
        actor_id = user_id or project.owner_id
        if project.status != models.ProcessingStatus.GENERATING:
            logger.info(f"[后台任务] 项目 {project_id} 当前不是生成状态，任务中止")
            return

        result = await db.execute(
            select(models.Scene)
            .where(models.Scene.project_id == project_id)
            .order_by(models.Scene.scene_index)
        )
        scenes = result.scalars().all()
        if not scenes:
            exc = RuntimeError("项目没有可生成的场次")
            project.status = models.ProcessingStatus.FAILED
            record_generation_error(project, exc, stage="content")
            await db.commit()
            logger.error(f"[后台任务] 项目 {project_id} 没有可生成的场次")
            return

        completed_scenes: list[models.Scene] = []
        story_bible = build_story_bible(
            logline=project.logline,
            project_type=project.project_type,
            genre=project.genre or "",
            global_context=project.global_context or {},
        )
        total_scenes = len(scenes)
        template_instructions = await get_prompt_addendum(
            db,
            stage="content",
            project_type=project.project_type,
        )

        for scene in scenes:
            try:
                await db.refresh(project)
            except Exception:
                logger.info(f"[后台任务] 项目 {project_id} 已不存在，任务中止")
                return

            if project.status != models.ProcessingStatus.GENERATING:
                logger.info("[后台任务] 检测到停止信号，任务中止")
                return

            await db.refresh(scene)
            if scene.status == models.ProcessingStatus.COMPLETED:
                completed_scenes.append(scene)
                continue

            scene_id = scene.id
            scene_index = scene.scene_index
            scene_outline = scene.outline
            logger.info(f"[后台任务] 正在生成第 {scene_index} 场: {scene_outline[:30]}...")
            scene.status = models.ProcessingStatus.GENERATING
            await db.commit()

            generated_content = ""
            usage = 0
            continuity_context = ""
            try:
                continuity_context = build_content_continuity_context(
                    story_bible=story_bible,
                    completed_scenes=completed_scenes,
                    current_index=scene_index,
                    total_scenes=total_scenes,
                )
                if project.project_type == "short_video":
                    generated_content, usage = await run_project_ai_call(
                        db=db, project=project, actor_id=actor_id, action=f"write_scene_{scene_index}",
                        prompt=continuity_context, expected_status=models.ProcessingStatus.GENERATING,
                        validate=validate_ai_text,
                        result_status=lambda value: "partial" if looks_like_story_restart(str(value), scene_index) else "success",
                        invoke=lambda: llm.write_short_video_prompt(
                            logline=project.logline,
                            style_guide=project.genre,
                            current_scene_outline=scene_outline,
                            clip_index=scene_index,
                            previous_context=continuity_context,
                            template_instructions=template_instructions,
                        ),
                    )
                else:
                    generated_content, usage = await run_project_ai_call(
                        db=db, project=project, actor_id=actor_id, action=f"write_scene_{scene_index}",
                        prompt=continuity_context, expected_status=models.ProcessingStatus.GENERATING,
                        validate=validate_ai_text,
                        result_status=lambda value: "partial" if looks_like_story_restart(str(value), scene_index) else "success",
                        invoke=lambda: llm.write_scene_content(
                            logline=project.logline,
                            style_guide=project.genre,
                            current_scene_outline=scene_outline,
                            previous_context=continuity_context,
                            scene_index=scene_index,
                            total_scenes=total_scenes,
                            template_instructions=template_instructions,
                        ),
                    )

                await db.refresh(project)
                if project.status != models.ProcessingStatus.GENERATING:
                    logger.info(
                        "[后台任务] 第 %s 场结果因任务已取消而丢弃。",
                        scene_index,
                    )
                    return

                generated_content = str(generated_content or "").strip()
                if not generated_content:
                    raise RuntimeError("LLM returned empty scene content")
                if looks_like_story_restart(generated_content, scene_index):
                    logger.warning(
                        "[后台任务] 第 %s 场正文疑似重新开篇，启用连续性守卫重写一次。",
                        scene_index,
                    )
                    guarded_context = continuity_context + (
                        "\n【连续性守卫】上一次正文疑似重新开篇。请直接承接上一场的最后动作，"
                        "禁止使用‘故事开始、序幕、初次相遇、第一次见面’等重启表达，"
                        "并保留人物当前的记忆、关系、位置、伤情、道具和未解线索。"
                    )
                    if project.project_type == "short_video":
                        retry_content, retry_usage = await run_project_ai_call(
                            db=db, project=project, actor_id=actor_id, action=f"write_scene_{scene_index}",
                            prompt=guarded_context, attempt=2, expected_status=models.ProcessingStatus.GENERATING,
                            validate=lambda value: validate_continuation(value, scene_index),
                            invoke=lambda: llm.write_short_video_prompt(
                                logline=project.logline,
                                style_guide=project.genre,
                                current_scene_outline=scene_outline,
                                clip_index=scene_index,
                                previous_context=guarded_context,
                                template_instructions=template_instructions,
                            ),
                        )
                    else:
                        retry_content, retry_usage = await run_project_ai_call(
                            db=db, project=project, actor_id=actor_id, action=f"write_scene_{scene_index}",
                            prompt=guarded_context, attempt=2, expected_status=models.ProcessingStatus.GENERATING,
                            validate=lambda value: validate_continuation(value, scene_index),
                            invoke=lambda: llm.write_scene_content(
                                logline=project.logline,
                                style_guide=project.genre,
                                current_scene_outline=scene_outline,
                                previous_context=guarded_context,
                                scene_index=scene_index,
                                total_scenes=total_scenes,
                                template_instructions=template_instructions,
                            ),
                        )
                    await db.refresh(project)
                    if project.status != models.ProcessingStatus.GENERATING:
                        logger.info(
                            "[后台任务] 第 %s 场重写结果因任务已取消而丢弃。",
                            scene_index,
                        )
                        return
                    usage += int(retry_usage or 0)
                    generated_content = str(retry_content or "").strip()
                    if not generated_content or looks_like_story_restart(
                        generated_content, scene_index
                    ):
                        raise RuntimeError("连续性守卫拒绝了疑似重新开篇的场次正文")

                scene.content = generated_content
                scene.summary = summarize_scene_for_continuity(
                    scene_outline,
                    generated_content,
                )
                scene.status = models.ProcessingStatus.COMPLETED
                completed_scenes.append(scene)
                project.global_summary = "\n".join(
                    f"第{item.scene_index}场：{item.summary or item.outline}"
                    for item in completed_scenes[-12:]
                )[:12000]
                await db.commit()
                logger.info(f"[后台任务] 第 {scene_index} 场生成完成")
            except Exception as exc:
                await db.rollback()
                if isinstance(exc, HTTPException) and exc.status_code == 409:
                    return
                project = await db.get(models.Project, project_id)
                failed_scene = await db.get(models.Scene, scene_id)
                if project:
                    project.status = models.ProcessingStatus.FAILED
                    record_generation_error(project, exc, stage=f"scene:{scene_index}")
                if failed_scene:
                    failed_scene.status = models.ProcessingStatus.FAILED
                await db.commit()

                logger.exception(
                    f"[后台任务] 第 {scene_index} 场生成失败: {exc}"
                )
                return

        project.status = models.ProcessingStatus.COMPLETED
        await db.commit()
        logger.info(f"[后台任务] 项目 {project_id} 所有剧本生成任务完成！")


async def run_generation_loop(project_id: int, user_id: Optional[int] = None):
    """Run the generation engine and guarantee a terminal failure state."""
    try:
        await _run_generation_loop(project_id, user_id=user_id)
    except Exception as exc:
        logger.exception(
            "[后台任务] 项目 %s 发生未处理的生成错误: %s",
            project_id,
            exc,
        )
        try:
            async with database.SessionLocal() as db:
                await mark_claimed_project_failed(db, project_id)
                project = await db.get(models.Project, project_id)
                if project:
                    record_generation_error(project, exc, stage="content")
                await db.execute(
                    update(models.Scene)
                    .where(models.Scene.project_id == project_id)
                    .where(models.Scene.status == models.ProcessingStatus.GENERATING)
                    .values(status=models.ProcessingStatus.FAILED)
                )
                await db.commit()
        except Exception as recovery_exc:
            logger.critical(
                "[后台任务] 项目 %s 的失败状态无法写入: %s",
                project_id,
                recovery_exc,
            )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "main:app",
        host="127.0.0.1",
        port=8000,
        reload=True,
        log_level="info",
    )
